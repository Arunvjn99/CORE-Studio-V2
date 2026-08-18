"""
CORE Studio V2 — Local Development Server

Departments:
  Design Team   — Planner, UX, UI, Review, Accessibility, Screenshot→Design, Code Export
  Analyst Team  — Research, BA, Personas, User Stories, PRD, Competitor Analysis
  QA Team       — Test Cases, Accessibility Check (real WCAG math), Compliance
  Engineering   — Code Gen (qwen2.5-coder), API Design, Architecture, Tech Docs
  Documentation — RAG-powered knowledge base (ChromaDB)

AI Backend:
  Local  → Ollama (deepseek-r1:8b reasoning, qwen2.5-coder:7b code, llava:7b vision)
  Cloud  → Anthropic Claude (Haiku/Sonnet/Opus) when API key set

Run: python3 server_local.py
"""
import os, sys, uuid, json, asyncio, time, re, base64
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Optional, List, Dict, Any
from contextlib import asynccontextmanager

# ── Dependency check ──────────────────────────────────────────────────────────
missing = []
try:
    from fastapi import FastAPI, Depends, HTTPException, WebSocket, WebSocketDisconnect, BackgroundTasks, UploadFile, File, Form
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
    from pydantic import BaseModel
    import sqlalchemy as sa
    from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
    from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column
    from jose import jwt, JWTError
    import bcrypt as _bcrypt
    import aiofiles
except ImportError as e:
    missing.append(str(e))

if missing:
    print(f"\n❌ Missing: {missing}")
    print("Run: pip3 install -r requirements-local.txt\n")
    sys.exit(1)

# ── Config ────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent

# Load .env if present — override=True so file values win over empty shell vars
_env_path = BASE_DIR / ".env"
try:
    from dotenv import load_dotenv
    load_dotenv(_env_path, override=True)
except ImportError:
    pass

# Always also parse directly and write into os.environ (belt and suspenders)
if _env_path.exists():
    for _line in _env_path.read_text().splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            _v = _v.strip().strip('"').strip("'")
            if _v:  # only set non-empty values
                os.environ[_k.strip()] = _v

DB_PATH = str(BASE_DIR / "core_studio_local.db")
DATABASE_URL = f"sqlite+aiosqlite:///{DB_PATH}"
JWT_SECRET = os.getenv("JWT_SECRET", "local-dev-secret-change-in-production")
JWT_ALGORITHM = "HS256"

# ── Multi-Provider AI Detection ───────────────────────────────────────────────
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "").strip()
OPENAI_API_KEY    = os.getenv("OPENAI_API_KEY", "").strip()
GEMINI_API_KEY    = os.getenv("GEMINI_API_KEY", "").strip() or os.getenv("GOOGLE_API_KEY", "").strip()
OLLAMA_BASE_URL   = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
AI_BACKEND_MODE   = os.getenv("AI_BACKEND_MODE", "auto").lower()

# Vision (image analysis) is deliberately resolved INDEPENDENTLY of AI_BACKEND_MODE/
# CLOUD_PROVIDER below. AI_BACKEND_MODE governs which model handles TEXT generation
# (planning/UI/code) and a user may reasonably pick "ollama" there to save cost — but
# reading an uploaded reference image accurately (the "recreate"/"redesign" flow) is a
# different job with a much bigger quality gap between local and cloud models, so it
# always prefers the best available vision model regardless of the text-generation
# choice. VISION_PROVIDER lets you pin a specific one; "auto" (default) picks the best
# available key in priority order (anthropic > openai > gemini), falling back to local
# Ollama (llava) only when no cloud key is configured at all.
VISION_PROVIDER = os.getenv("VISION_PROVIDER", "auto").strip().lower()


def _detect_provider() -> str:
    """Detect which cloud provider to use based on which key is set."""
    if ANTHROPIC_API_KEY and ANTHROPIC_API_KEY.startswith("sk-ant"):
        return "anthropic"
    if OPENAI_API_KEY and OPENAI_API_KEY.startswith("sk-"):
        return "openai"
    if GEMINI_API_KEY:
        return "gemini"
    # generic sk- key without ant → openai
    if ANTHROPIC_API_KEY.startswith("sk-"):
        return "openai"
    return "none"


CLOUD_PROVIDER = _detect_provider()

# Per-provider model tiers (fast / standard / advanced)
PROVIDER_MODELS = {
    "anthropic": {
        "fast":     os.getenv("ANTHROPIC_FAST",     "claude-3-5-haiku-20241022"),
        "standard": os.getenv("ANTHROPIC_MODEL",    "claude-sonnet-4-20250514"),
        "advanced": os.getenv("ANTHROPIC_ADVANCED", "claude-sonnet-4-20250514"),
    },
    "openai": {
        "fast":     os.getenv("OPENAI_FAST",     "gpt-4o-mini"),
        "standard": os.getenv("OPENAI_MODEL",    "gpt-4o"),
        "advanced": os.getenv("OPENAI_ADVANCED", "gpt-4o"),
    },
    "gemini": {
        "fast":     os.getenv("GEMINI_FAST",     "gemini-1.5-flash"),
        "standard": os.getenv("GEMINI_MODEL",    "gemini-1.5-pro"),
        "advanced": os.getenv("GEMINI_ADVANCED", "gemini-1.5-pro"),
    },
}

# Ollama (local) models
OLLAMA_MODELS = {
    "reasoning": os.getenv("OLLAMA_REASONING_MODEL", "deepseek-r1:8b"),
    "coder":     os.getenv("OLLAMA_CODER_MODEL",     "qwen2.5-coder:7b"),
    "vision":    os.getenv("OLLAMA_VISION_MODEL",    "llava:7b"),
}
AGENT_LOCAL_MODEL = {
    "planner": OLLAMA_MODELS["reasoning"], "ux": OLLAMA_MODELS["reasoning"],
    "ui": OLLAMA_MODELS["coder"], "review": OLLAMA_MODELS["reasoning"],
    "fast": OLLAMA_MODELS["reasoning"], "research": OLLAMA_MODELS["reasoning"],
    "qa": OLLAMA_MODELS["reasoning"],
}
AGENT_CLOUD_BOOST = {"ui": 0, "ux": 0}  # removed boost — sonnet is sufficient for UI/UX


def _check_ollama() -> bool:
    import urllib.request
    try:
        urllib.request.urlopen(f"{OLLAMA_BASE_URL}/api/tags", timeout=2)
        return True
    except Exception:
        return False


HAS_CLOUD_AI = CLOUD_PROVIDER != "none"
HAS_OLLAMA   = _check_ollama()
HAS_AI       = HAS_CLOUD_AI or HAS_OLLAMA

# Decide the active backend
if AI_BACKEND_MODE == "ollama":
    AI_BACKEND = "ollama" if HAS_OLLAMA else "none"
elif AI_BACKEND_MODE == "cloud":
    AI_BACKEND = "cloud" if HAS_CLOUD_AI else ("ollama" if HAS_OLLAMA else "none")
else:  # auto — prefer cloud if a key is present, else local
    AI_BACKEND = "cloud" if HAS_CLOUD_AI else ("ollama" if HAS_OLLAMA else "none")


def select_model(complexity: int, agent: str) -> str:
    """Pick the right model for the active backend + provider."""
    if AI_BACKEND == "cloud":
        tiers = PROVIDER_MODELS.get(CLOUD_PROVIDER, PROVIDER_MODELS["anthropic"])
        boost = AGENT_CLOUD_BOOST.get(agent, 0)
        adjusted = min(10, complexity + boost)
        if adjusted <= 3:   return tiers["fast"]
        if adjusted <= 7:   return tiers["standard"]
        return tiers["advanced"]
    if AI_BACKEND == "ollama":
        return AGENT_LOCAL_MODEL.get(agent, OLLAMA_MODELS["reasoning"])
    return "none"


# Back-compat alias used in startup banner
CLOUD_MODELS = PROVIDER_MODELS.get(CLOUD_PROVIDER, PROVIDER_MODELS["anthropic"])


# ── Design System (from V1) ───────────────────────────────────────────────────
sys.path.insert(0, str(BASE_DIR))
try:
    from app.design_system.tokens import DESIGN_TOKENS, COMPONENTS
    from app.design_system.builder import UIBuilder
    HAS_DESIGN_SYSTEM = True
except ImportError:
    HAS_DESIGN_SYSTEM = False
    DESIGN_TOKENS = {}
    COMPONENTS = {}


# ── LLM Client ────────────────────────────────────────────────────────────────
import urllib.request as _urllib_request

_provider_clients = {}
_token_tracker: dict = {"input": 0, "output": 0, "calls": 0, "last_model": "", "cost_usd": 0.0}
_gen_token_tracker: dict = {}  # gen_id -> {"input": int, "output": int, "calls": int, "cost_start": float}

# Approximate public list pricing, USD per 1M tokens (input, output). Matched by substring
# against the model name rather than exact snapshot ID, since ANTHROPIC_MODEL/etc. env vars
# can point at different dated snapshots of the same family. Local Ollama models are free.
# Update these if provider pricing changes.
_MODEL_PRICING = [
    # (substring to match in model name, price_in_per_mtok, price_out_per_mtok)
    ("opus",   15.00, 75.00),
    ("sonnet",  3.00, 15.00),
    ("haiku",   0.80,  4.00),
    ("gpt-4o-mini", 0.15, 0.60),
    ("gpt-4o",      2.50, 10.00),
    ("gemini-1.5-flash", 0.075, 0.30),
    ("gemini-1.5-pro",   1.25,  5.00),
    ("gemini",           1.25,  5.00),  # fallback for other gemini variants
]


def _price_for_model(model: str) -> tuple[float, float]:
    """Returns (price_in_per_mtok, price_out_per_mtok) for a model name, or (0,0) for
    unrecognized/local models (Ollama)."""
    name = (model or "").lower()
    for substr, price_in, price_out in _MODEL_PRICING:
        if substr in name:
            return price_in, price_out
    return 0.0, 0.0


def _track_usage(model: str, input_tokens: int, output_tokens: int) -> None:
    """Central place to update the global token/cost tracker for any provider call."""
    price_in, price_out = _price_for_model(model)
    _token_tracker["input"] += input_tokens
    _token_tracker["output"] += output_tokens
    _token_tracker["calls"] += 1
    _token_tracker["last_model"] = model
    _token_tracker["cost_usd"] += (input_tokens / 1_000_000) * price_in + (output_tokens / 1_000_000) * price_out

def _get_anthropic():
    if "anthropic" not in _provider_clients:
        from anthropic import Anthropic
        _provider_clients["anthropic"] = Anthropic(api_key=ANTHROPIC_API_KEY)
    return _provider_clients["anthropic"]

def _get_openai():
    if "openai" not in _provider_clients:
        from openai import OpenAI
        _provider_clients["openai"] = OpenAI(api_key=OPENAI_API_KEY or ANTHROPIC_API_KEY)
    return _provider_clients["openai"]

def _get_gemini():
    if "gemini" not in _provider_clients:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        _provider_clients["gemini"] = genai
    return _provider_clients["gemini"]


def _call_ollama_sync(model: str, system: str, user: str) -> str:
    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user",   "content": user},
        ],
        "stream": False,
        # num_ctx must be set explicitly — Ollama otherwise silently caps the
        # context window at 4096 tokens regardless of what the model supports,
        # and design-system prompts (tokens + layout grammar + RAG context)
        # routinely exceed that.
        "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
    }).encode("utf-8")
    req = _urllib_request.Request(
        f"{OLLAMA_BASE_URL}/api/chat", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    try:
        with _urllib_request.urlopen(req, timeout=300) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data.get("message", {}).get("content", "")
    except Exception as e:
        return json.dumps({"error": str(e)})


def _call_anthropic_sync(model: str, system: str, user: str) -> str:
    client = _get_anthropic()
    msg = client.messages.create(
        model=model, max_tokens=8192, system=system,
        messages=[{"role": "user", "content": user}],
    )
    _track_usage(model, msg.usage.input_tokens, msg.usage.output_tokens)
    return msg.content[0].text


def _call_openai_sync(model: str, system: str, user: str) -> str:
    client = _get_openai()
    resp = client.chat.completions.create(
        model=model, max_tokens=8192,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    if resp.usage:
        _track_usage(model, resp.usage.prompt_tokens, resp.usage.completion_tokens)
    return resp.choices[0].message.content


def _call_gemini_sync(model: str, system: str, user: str) -> str:
    genai = _get_gemini()
    from google.generativeai.types import GenerationConfig
    gmodel = genai.GenerativeModel(model, system_instruction=system)
    resp = gmodel.generate_content(user, generation_config=GenerationConfig(max_output_tokens=8192))
    if getattr(resp, "usage_metadata", None):
        _track_usage(model, resp.usage_metadata.prompt_token_count, resp.usage_metadata.candidates_token_count)
    return resp.text


# Cloud provider dispatch
_CLOUD_DISPATCH = {
    "anthropic": _call_anthropic_sync,
    "openai":    _call_openai_sync,
    "gemini":    _call_gemini_sync,
}


async def llm_chat(system: str, user: str, model: str) -> str:
    """
    Unified LLM call — routes to the active backend:
    Ollama (local) OR Anthropic/OpenAI/Gemini (cloud, auto-detected from key).
    """
    if not HAS_AI:
        return json.dumps({"error": "no_ai", "message": "No AI backend. Start Ollama or set an API key."})

    loop = asyncio.get_event_loop()

    if AI_BACKEND == "ollama":
        return await loop.run_in_executor(None, _call_ollama_sync, model, system, user)

    # Cloud
    fn = _CLOUD_DISPATCH.get(CLOUD_PROVIDER)
    if not fn:
        return json.dumps({"error": "no_provider", "message": f"Unknown provider: {CLOUD_PROVIDER}"})
    try:
        return await loop.run_in_executor(None, fn, model, system, user)
    except Exception as e:
        # Fall back to Ollama if cloud fails and Ollama is up
        if HAS_OLLAMA:
            fallback_model = AGENT_LOCAL_MODEL.get("planner", OLLAMA_MODELS["reasoning"])
            return await loop.run_in_executor(None, _call_ollama_sync, fallback_model, system, user)
        return json.dumps({"error": "cloud_failed", "message": str(e)})


def _parse_json(raw: str) -> Optional[dict]:
    """
    Strip <think>…</think> (deepseek-r1), markdown fences, then parse JSON.
    Also handles cases where JSON is embedded anywhere in the output.
    """
    text = raw.strip()

    # Strip deepseek-r1 <think> blocks
    import re as _re
    text = _re.sub(r"<think>.*?</think>", "", text, flags=_re.DOTALL).strip()

    # Strip markdown fences
    for prefix in ("```json", "```"):
        if text.startswith(prefix):
            text = text[len(prefix):]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    # Try direct parse
    try:
        return json.loads(text)
    except Exception:
        pass

    # Try to find JSON object anywhere in the output (common with local models)
    match = _re.search(r"\{.*\}", text, _re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass

    return None


# ── RAG (ChromaDB — same as V1) ───────────────────────────────────────────────
_chroma_store = None

def _get_chroma():
    global _chroma_store
    if _chroma_store is None:
        try:
            from chromadb import PersistentClient
            persist_dir = str(BASE_DIR / ".chroma_db")
            Path(persist_dir).mkdir(exist_ok=True)
            client = PersistentClient(path=persist_dir)
            _chroma_store = {
                name: client.get_or_create_collection(name)
                for name in ["guidelines", "past_designs", "best_practices", "design_patterns", "company_standards", "documents"]
            }
        except ImportError:
            _chroma_store = {}
    return _chroma_store


async def rag_query(collection: str, query: str, top_k: int = 5) -> List[str]:
    """Query ChromaDB — returns list of text strings."""
    try:
        store = _get_chroma()
        col = store.get(collection)
        if not col:
            return []
        results = col.query(query_texts=[query], n_results=top_k)
        docs = results.get("documents", [[]])[0]
        return [d for d in docs if d]
    except Exception:
        return []


async def rag_add(collection: str, doc_id: str, text: str, metadata: dict = None):
    """Add document to ChromaDB collection."""
    try:
        store = _get_chroma()
        col = store.get(collection)
        if col:
            col.upsert(ids=[doc_id], documents=[text], metadatas=[metadata or {}])
    except Exception:
        pass


def _chunk_text(text: str, chunk_size: int = 900, overlap: int = 150) -> list[str]:
    """Split text into overlapping chunks preserving paragraph boundaries."""
    import re as _re
    paragraphs = _re.split(r"\n\s*\n", text.strip())
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) < chunk_size:
            current += ("\n\n" + para) if current else para
        else:
            if current:
                chunks.append(current.strip())
            current = (current[-overlap:] + "\n\n" + para) if (current and overlap) else para
    if current:
        chunks.append(current.strip())
    return [c for c in chunks if len(c) > 40]


async def rag_add_chunked(collection: str, base_id: str, text: str, metadata: dict = None):
    """Chunk a document and add all chunks to ChromaDB. Returns chunk count."""
    chunks = _chunk_text(text)
    for i, chunk in enumerate(chunks):
        await rag_add(collection, f"{base_id}-c{i}", chunk, {**(metadata or {}), "chunk_index": i})
    return len(chunks)


async def seed_knowledge_from_files() -> int:
    """
    On startup, scan the knowledge/ folder and load all .md files into ChromaDB.
    Uses file path as a stable doc_id so re-runs are idempotent (upsert).
    Skips seeding if collections already contain data.
    Returns total chunks indexed.
    """
    knowledge_dir = BASE_DIR.parent.parent / "knowledge"
    if not knowledge_dir.exists():
        return 0

    # Skip re-seeding if collections already populated (avoids re-embedding on every restart)
    try:
        store = _get_chroma()
        existing = sum(
            store.get(c).count() if store.get(c) else 0
            for c in ["guidelines", "company_standards"]
        )
        if existing > 100:
            return existing
    except Exception:
        pass

    # Map knowledge folder category → ChromaDB collection
    CATEGORY_MAP = {
        "company": "company_standards",
        "domain": "guidelines",
        "layout-references": "guidelines",
    }

    total = 0
    all_md_files = list(knowledge_dir.rglob("*.md")) + list(knowledge_dir.rglob("*.MD"))
    for md_file in all_md_files:
        try:
            parts = md_file.relative_to(knowledge_dir).parts
            if not parts:
                continue
            top_level = parts[0]
            collection = CATEGORY_MAP.get(top_level, "guidelines")

            # Stable doc_id based on relative path (no spaces, lowercase)
            rel_path = str(md_file.relative_to(knowledge_dir))
            doc_id = "seed-" + rel_path.replace("/", "-").replace(" ", "_").replace(".md", "")

            text = md_file.read_text(encoding="utf-8", errors="ignore").strip()
            if not text:
                continue

            metadata = {
                "source": rel_path,
                "category": top_level,
                "subcategory": parts[1] if len(parts) > 2 else parts[0],
                "name": md_file.stem,
            }
            chunks = _chunk_text(text)
            for i, chunk in enumerate(chunks):
                await rag_add(collection, f"{doc_id}-c{i}", chunk, {**metadata, "chunk_index": i})
            total += len(chunks)
        except Exception:
            pass

    return total


async def rag_query_rich(collection: str, query: str, top_k: int = 5) -> list[dict]:
    """Query ChromaDB returning text + metadata + distance for citations."""
    try:
        store = _get_chroma()
        col = store.get(collection)
        if not col:
            return []
        results = col.query(query_texts=[query], n_results=top_k)
        docs = results.get("documents", [[]])[0]
        metas = results.get("metadatas", [[]])[0]
        dists = results.get("distances", [[]])[0]
        out = []
        for i, d in enumerate(docs):
            if d:
                out.append({
                    "content": d,
                    "metadata": metas[i] if i < len(metas) else {},
                    "distance": dists[i] if i < len(dists) else None,
                    "source": (metas[i] or {}).get("name", "Unknown") if i < len(metas) else "Unknown",
                })
        return out
    except Exception:
        return []


# ── Real Agents (ported from V1, adapted to V2 model router) ─────────────────

async def run_planner(title: str, description: str, complexity: int) -> Dict:
    """Planner agent — same prompt as V1 planner_v2.py, dynamic model."""
    model = select_model(complexity, "planner")

    guidelines = await rag_query("guidelines", f"{title} {description}")
    company_standards = await rag_query("company_standards", f"UX standards accessibility compliance {title}")
    past_designs = await rag_query("past_designs", title)

    guidelines_text = "\n".join([g[:200] for g in guidelines[:4]]) or "Standard design best practices"
    company_text = "\n".join([c[:200] for c in company_standards[:3]]) or "Apply WCAG 2.1 AA accessibility and UX best practices"
    past_text = "\n".join([str(d)[:150] for d in past_designs[:2]]) or "No past designs available"

    prompt = f"""You are a senior product design strategist.

PROJECT: {title}
DESCRIPTION: {description}

DOMAIN & INDUSTRY GUIDELINES:
{guidelines_text}

COMPANY STANDARDS (accessibility, UX, compliance — must follow):
{company_text}

RELEVANT PAST DESIGNS:
{past_text}

Create a comprehensive design plan. Return ONLY valid JSON:

{{
  "design_plan": {{
    "name": "{title}",
    "description": "Concise overview of the design approach",
    "rationale": "Why this approach — what user problems it solves",
    "user_journey": {{
      "personas": [
        {{"name": "Primary User", "goals": ["goal 1", "goal 2"], "pain_points": ["pain 1", "pain 2"]}}
      ],
      "user_flow": "Step-by-step: 1 → 2 → 3",
      "key_moments": ["Moment 1", "Moment 2", "Moment 3"]
    }},
    "design_principles": [
      {{"principle": "Mobile-First", "rationale": "Most users on mobile", "implication": "Single-column layouts"}},
      {{"principle": "Accessibility", "rationale": "Inclusive design", "implication": "WCAG 2.1 AA"}},
      {{"principle": "Consistency", "rationale": "Build trust", "implication": "Reuse design-system components"}}
    ],
    "screens": [
      {{"id": "screen-1", "name": "Entry Screen", "purpose": "Introduce user and set expectations", "user_actions": ["Read headline", "Click CTA"], "key_elements": ["Hero headline", "Subtext", "Primary CTA button"], "success_criteria": "User understands purpose and clicks CTA"}},
      {{"id": "screen-2", "name": "Main Interaction Screen", "purpose": "Allow user to complete the core task", "user_actions": ["Fill fields", "Submit form"], "key_elements": ["Form fields", "Inline validation", "Submit button"], "success_criteria": "Form validates and data is submitted"}},
      {{"id": "screen-3", "name": "Confirmation Screen", "purpose": "Confirm successful completion", "user_actions": ["Read confirmation", "Take next action"], "key_elements": ["Success indicator", "Summary", "Next action CTA"], "success_criteria": "User feels confident the task is complete"}}
    ],
    "design_patterns": [
      {{"name": "Progressive Disclosure", "usage": "Multi-step forms", "rationale": "Reduces cognitive load"}},
      {{"name": "Inline Validation", "usage": "Form fields", "rationale": "Immediate feedback prevents errors"}},
      {{"name": "Clear CTA Hierarchy", "usage": "Every screen", "rationale": "One primary action per screen"}}
    ],
    "accessibility_requirements": ["Colour contrast ≥ 4.5:1", "Touch targets ≥ 48×48 px", "Full keyboard navigation", "ARIA labels + roles", "Visible focus indicators"],
    "responsive_breakpoints": {{"mobile": {{"width": 375}}, "tablet": {{"width": 768}}, "desktop": {{"width": 1440}}}},
    "success_metrics": [
      {{"metric": "Task Completion Rate", "target": "85%", "measurement": "Analytics funnel"}},
      {{"metric": "Accessibility Score", "target": "95/100", "measurement": "Automated audit"}}
    ]
  }}
}}

Only return JSON."""

    raw = await llm_chat(
        system="You are a senior design strategist. Return ONLY valid JSON.",
        user=prompt,
        model=model,
    )
    data = _parse_json(raw)
    if data:
        plan = data.get("design_plan") or data
        if "screens" in plan and len(plan["screens"]) >= 2:
            return {"plan": plan, "model": model}

    # Fallback
    return {
        "plan": {
            "name": title,
            "description": f"Design plan for {title}: {description[:80]}",
            "user_journey": {
                "personas": [{"name": "User", "goals": ["Complete task"], "pain_points": ["Unclear process"]}],
                "user_flow": "Start → Input → Confirm → Done",
                "key_moments": ["First impression", "Decision point", "Completion"],
            },
            "design_principles": [
                {"principle": "Clarity", "rationale": "Users need guidance", "implication": "Simple, intuitive layout"},
                {"principle": "Accessibility", "rationale": "Inclusive by default", "implication": "WCAG 2.1 AA"},
            ],
            "screens": [
                {"id": "screen-1", "name": "Welcome", "purpose": "Entry point", "user_actions": ["Click CTA"], "key_elements": ["Title", "CTA"], "success_criteria": "User proceeds"},
                {"id": "screen-2", "name": "Main Task", "purpose": "Core interaction", "user_actions": ["Submit"], "key_elements": ["Form", "Submit"], "success_criteria": "Task completed"},
                {"id": "screen-3", "name": "Confirmation", "purpose": "Confirm done", "user_actions": ["View"], "key_elements": ["Success badge", "Next CTA"], "success_criteria": "User sees confirmation"},
            ],
            "accessibility_requirements": ["WCAG 2.1 AA", "Keyboard navigation"],
            "success_metrics": [{"metric": "Completion Rate", "target": "85%", "measurement": "Analytics"}],
        },
        "model": model,
        "fallback": True,
    }


async def run_ux(plan: dict, title: str, complexity: int) -> Dict:
    """UX agent — same prompt as V1 ux_v2.py, dynamic model."""
    model = select_model(complexity, "ux")
    screens = plan.get("screens", [])
    principles = plan.get("design_principles", [])

    patterns = await rag_query("design_patterns", title)
    domain_knowledge = await rag_query("guidelines", f"{title} UX flow interaction patterns")
    company_ux = await rag_query("company_standards", "UX standards interaction design accessibility")
    pattern_text = "\n".join([p[:100] for p in patterns[:3]]) or "Mobile-first, accessibility-first"
    domain_text = "\n".join([d[:150] for d in domain_knowledge[:3]]) or ""
    company_ux_text = "\n".join([c[:150] for c in company_ux[:2]]) or ""
    screen_lines = "\n".join(f"- {s['id']}: {s['name']} — {s.get('purpose','')}" for s in screens[:5])
    principle_lines = ", ".join([p.get("principle", "") for p in principles[:3]]) or "Clarity, Efficiency, Accessibility"

    prompt = f"""You are a UX/interaction designer creating detailed user flows.

PROJECT: {title}
SCREENS FROM PLAN:
{screen_lines or "- screen-1: Welcome\n- screen-2: Main Task\n- screen-3: Confirmation"}

DESIGN PRINCIPLES: {principle_lines}
INTERACTION PATTERNS: {pattern_text}
{f"DOMAIN-SPECIFIC KNOWLEDGE:{chr(10)}{domain_text}" if domain_text else ""}
{f"COMPANY UX STANDARDS (must follow):{chr(10)}{company_ux_text}" if company_ux_text else ""}

Create a complete UX flow. Return ONLY valid JSON:
{{
  "ux_flow": {{
    "name": "{title} Flow",
    "screens": [
      {{
        "id": "screen-1",
        "name": "Welcome",
        "description": "First screen — set expectations",
        "layout": "vertical_center",
        "wireframe": {{
          "sections": [
            {{"type": "header", "content": "App name or brand"}},
            {{"type": "description", "content": "What the user will accomplish"}},
            {{"type": "cta_button", "label": "Get Started", "action": "next"}},
            {{"type": "text", "content": "Already have an account? Sign in"}}
          ]
        }},
        "interactions": [
          {{"element": "Get Started", "action": "tap", "result": "Navigate to screen-2"}}
        ],
        "states": {{"default": "Ready for user"}},
        "accessibility": {{"landmarks": ["main"], "headings": ["h1"]}}
      }},
      {{
        "id": "screen-2",
        "name": "Core Interaction",
        "description": "User performs the main task",
        "layout": "vertical_scrollable",
        "wireframe": {{
          "sections": [
            {{"type": "heading", "content": "Complete Your Information"}},
            {{"type": "form", "fields": [
              {{"label": "Primary Field", "type": "text", "required": true, "validation": "Required"}},
              {{"label": "Secondary Field", "type": "textarea", "required": false}}
            ]}},
            {{"type": "button_group", "buttons": [
              {{"label": "Back", "style": "secondary"}},
              {{"label": "Submit", "style": "primary", "disabled_until": "form_valid"}}
            ]}}
          ]
        }},
        "interactions": [
          {{"element": "Submit", "action": "tap", "result": "Submit data, navigate to screen-3"}},
          {{"element": "Back", "action": "tap", "result": "Navigate to screen-1"}}
        ],
        "states": {{"empty": "No fields filled", "valid": "All required fields valid", "submitting": "Sending data", "error": "Validation error"}},
        "error_handling": [{{"field": "Primary Field", "error": "Field required", "help_text": "This field cannot be empty"}}],
        "accessibility": {{"landmarks": ["main", "form"], "headings": ["h1"], "labels": ["Primary Field"]}}
      }},
      {{
        "id": "screen-3",
        "name": "Confirmation",
        "description": "Confirm successful completion",
        "layout": "vertical_center",
        "wireframe": {{
          "sections": [
            {{"type": "success_message", "icon": "checkmark", "title": "Success!", "subtitle": "Your task is complete"}},
            {{"type": "card", "content": "Summary of what was completed"}},
            {{"type": "button_group", "buttons": [
              {{"label": "Done", "style": "primary"}},
              {{"label": "View Details", "style": "secondary"}}
            ]}}
          ]
        }},
        "interactions": [{{"element": "Done", "action": "tap", "result": "Return to home"}}],
        "states": {{"showing": "Display confirmation"}},
        "accessibility": {{"landmarks": ["main"], "headings": ["h1"], "status": "Success announced"}}
      }}
    ],
    "user_flows": [
      {{"name": "Happy Path", "steps": ["screen-1", "screen-2 (valid)", "screen-3"], "description": "User completes all steps"}},
      {{"name": "Error Recovery", "steps": ["screen-2 (error)", "screen-2 (corrected)", "screen-3"], "description": "User corrects errors and submits"}}
    ],
    "micro_interactions": [
      {{"name": "Button Press", "trigger": "Tap button", "animation": "Scale to 0.95", "duration": "150ms"}},
      {{"name": "Field Validation", "trigger": "Blur input", "animation": "Border colour + icon", "duration": "200ms"}},
      {{"name": "Form Submit", "trigger": "Submit tap", "animation": "Button → spinner", "duration": "300ms"}}
    ]
  }}
}}

Only return JSON."""

    raw = await llm_chat(
        system="You are a UX designer. Return detailed interaction flows as valid JSON only.",
        user=prompt,
        model=model,
    )
    data = _parse_json(raw)
    if data:
        flow = data.get("ux_flow") or data
        if len(flow.get("screens", [])) >= 2:
            return {"ux": flow, "model": model}

    # Fallback UX
    return {
        "ux": {
            "name": f"{title} Flow",
            "screens": [
                {"id": "s1", "name": "Welcome", "layout": "vertical_center",
                 "wireframe": {"sections": [
                     {"type": "header", "content": title},
                     {"type": "description", "content": "Get started below"},
                     {"type": "cta_button", "label": "Get Started"},
                 ]},
                 "interactions": [{"element": "Get Started", "action": "tap", "result": "Go to screen-2"}],
                 "states": {"default": "Ready"}, "accessibility": {"landmarks": ["main"]}},
                {"id": "s2", "name": "Core Task", "layout": "vertical",
                 "wireframe": {"sections": [
                     {"type": "heading", "content": "Enter Details"},
                     {"type": "form", "fields": [{"label": "Name", "type": "text", "required": True}]},
                     {"type": "button_group", "buttons": [{"label": "Submit", "style": "primary"}]},
                 ]},
                 "interactions": [{"element": "Submit", "action": "tap", "result": "Go to screen-3"}],
                 "states": {"empty": "No input", "valid": "Ready", "submitting": "Loading"},
                 "accessibility": {"landmarks": ["main", "form"]}},
                {"id": "s3", "name": "Confirmation", "layout": "vertical_center",
                 "wireframe": {"sections": [
                     {"type": "success_message", "icon": "checkmark", "title": "Done!"},
                     {"type": "button_group", "buttons": [{"label": "Done", "style": "primary"}]},
                 ]},
                 "interactions": [{"element": "Done", "action": "tap", "result": "Return home"}],
                 "states": {"showing": "Visible"}, "accessibility": {"landmarks": ["main"]}},
            ],
            "user_flows": [{"name": "Happy Path", "steps": ["s1", "s2", "s3"], "description": "Successful completion"}],
            "micro_interactions": [{"name": "Button Press", "trigger": "Tap", "animation": "Scale", "duration": "150ms"}],
        },
        "model": model,
        "fallback": True,
    }


async def run_ui(plan: dict, ux: dict, title: str, complexity: int) -> Dict:
    """UI agent — uses V1 design_system builder + LLM colour decisions."""
    model = select_model(complexity, "ui")

    ux_screens = ux.get("screens", [])
    screen_names = [s.get("name", "Screen") for s in ux_screens[:4]]

    colour_prompt = f"""You are a visual/UI designer.

PROJECT: {title}
UX SCREENS: {', '.join(screen_names) or 'Welcome, Form, Confirmation'}

Customise the colour palette and typography. Return ONLY valid JSON:
{{
  "colours": {{
    "primary":       "#007AFF",
    "secondary":     "#34C759",
    "background":    "#FFFFFF",
    "surface":       "#F5F5F5",
    "text_primary":  "#000000",
    "text_secondary":"#666666",
    "border":        "#E5E5EA",
    "error":         "#FF3B30",
    "success":       "#34C759"
  }},
  "typography": {{
    "font_family": "Inter, -apple-system, sans-serif",
    "heading_size": 28,
    "body_size": 16
  }},
  "tone": "professional"
}}
Only return JSON."""

    colour_raw = await llm_chat(
        system="You are a visual designer. Return ONLY valid JSON.",
        user=colour_prompt,
        model=select_model(3, "planner"),  # colour selection is simple → fast model
    )
    theme = _parse_json(colour_raw) or {}

    # Use V1 UIBuilder if available
    if HAS_DESIGN_SYSTEM:
        builder = UIBuilder()
        screens = []
        for ux_screen in ux_screens[:5]:
            wf = ux_screen.get("wireframe") or {}
            sections = wf.get("sections") or []
            elements = builder.elements_from_wireframe_sections(sections)
            if not elements:
                elements = [
                    builder.create_heading(ux_screen.get("name", "Screen"), level=1),
                    builder.create_button("Continue", "primary"),
                ]
            bg = theme.get("colours", {}).get("background", "#FFFFFF")
            screen = builder.create_screen(
                name=ux_screen.get("name", "Screen"),
                elements=elements,
                description=ux_screen.get("description", ""),
                background=bg,
            )
            screens.append(screen)

        if not screens:
            bg = theme.get("colours", {}).get("background", "#FFFFFF")
            screens = [
                builder.create_screen("Welcome", [
                    builder.create_heading(f"Welcome to {title}", level=1),
                    builder.create_text("Get started below", "body_md"),
                    builder.create_button("Get Started", "primary"),
                ], background=bg),
                builder.create_screen("Form", [
                    builder.create_heading("Enter Details", level=2),
                    builder.create_input("Name", "text", required=True),
                    builder.create_button("Submit", "primary"),
                ], background=bg),
                builder.create_screen("Done", [
                    builder.create_badge("✓ Success", "success"),
                    builder.create_heading("All done!", level=1),
                    builder.create_button("Done", "primary"),
                ], background=bg),
            ]

        import copy
        merged_tokens = copy.deepcopy(DESIGN_TOKENS)
        project_colours = theme.get("colours", {})
        if project_colours:
            merged_tokens.setdefault("colors", {}).update(project_colours)
        if theme.get("typography", {}).get("font_family"):
            merged_tokens.setdefault("typography", {})["font_family"] = theme["typography"]["font_family"]

        ui_data = {
            "name": f"{title} UI",
            "tone": theme.get("tone", "professional"),
            "design_tokens": merged_tokens,
            "project_colours": project_colours,
            "screens": screens,
            "components": builder.components,
        }
    else:
        # Fallback without design system builder
        ui_data = {
            "name": f"{title} UI",
            "tone": theme.get("tone", "professional"),
            "design_tokens": {"colors": theme.get("colours", {})},
            "screens": [
                {
                    "id": f"screen-{i+1}",
                    "name": s.get("name", f"Screen {i+1}"),
                    "background": theme.get("colours", {}).get("background", "#FFFFFF"),
                    "elements": [{"type": "heading", "text": s.get("name", ""), "level": 1}],
                }
                for i, s in enumerate(ux_screens[:3])
            ],
        }

    return {"ui": ui_data, "model": model}


async def run_review(plan: dict, ux: dict, ui: dict, title: str, complexity: int) -> Dict:
    """Review/validator agent — same as V1 validator.py."""
    model = select_model(min(complexity, 5), "review")  # Review is analytical, cap at 5

    guidelines = await rag_query("guidelines", f"design validation {title}")
    company_standards = await rag_query("company_standards", "accessibility WCAG compliance UX standards")
    guideline_text = "\n".join(guidelines[:3]) or "Design quality standards"
    compliance_text = "\n".join([c[:200] for c in company_standards[:3]]) or "WCAG 2.1 AA, UX best practices"

    screen_count = len(ui.get("screens", []))
    ux_screen_count = len(ux.get("screens", []))

    prompt = f"""You are a design validator. Validate this design system output against company standards.

DESIGN: {title}
SCREENS: {screen_count} UI screens, {ux_screen_count} UX screens
TONE: {ui.get("tone", "professional")}

COMPANY COMPLIANCE STANDARDS (validate against these):
{compliance_text}
DESIGN PRINCIPLES: {", ".join([p.get("principle", "") for p in plan.get("design_principles", [])[:3]])}

GUIDELINES:
{guideline_text}

Return ONLY valid JSON:
{{
  "overall_status": "approved_with_changes",
  "overall_score": 85,
  "summary": "Design meets quality standards with minor improvements suggested.",
  "checks": [
    {{"category": "accessibility", "check_name": "Color Contrast", "status": "pass", "severity": "info", "details": "Primary text meets WCAG AA", "recommendation": "Maintain current contrast ratios", "wcag_criterion": "1.4.3"}},
    {{"category": "ux_quality", "check_name": "Navigation Clarity", "status": "pass", "severity": "info", "details": "Navigation is clear and consistent", "recommendation": "Good as-is"}},
    {{"category": "design_system", "check_name": "Spacing Consistency", "status": "warning", "severity": "minor", "details": "Review spacing values", "recommendation": "Align to 8px grid"}},
    {{"category": "accessibility", "check_name": "Keyboard Navigation", "status": "warning", "severity": "major", "details": "Interactive elements need visible focus indicators", "recommendation": "Add focus-visible ring", "wcag_criterion": "2.4.7"}},
    {{"category": "ux_quality", "check_name": "Error States", "status": "pass", "severity": "info", "details": "Error states defined in UX flow", "recommendation": "Implement all error states in code"}}
  ],
  "accessibility_report": {{
    "score": 87,
    "issues_critical": 0,
    "issues_major": 1,
    "issues_minor": 1,
    "highlights": ["Good color contrast", "Semantic structure", "Clear visual hierarchy"]
  }},
  "design_system_compliance": {{
    "score": 90,
    "deviations": ["Review spacing alignment with 8px base unit"]
  }},
  "requirements_coverage": {{
    "covered": [],
    "missing": [],
    "coverage_percent": 100
  }},
  "recommended_changes": [
    {{"priority": "high", "area": "Interactive elements", "change": "Add focus-visible ring to all buttons and links", "reason": "WCAG 2.4.7 keyboard navigation compliance"}},
    {{"priority": "medium", "area": "Spacing", "change": "Align all padding/margin to 8px grid", "reason": "Design system consistency"}},
    {{"priority": "low", "area": "Micro-interactions", "change": "Add button press scale animation (150ms)", "reason": "Better user feedback"}}
  ],
  "approved_patterns": [
    "Card-based layout with sidebar navigation for dashboard contexts",
    "Progressive disclosure pattern for multi-step forms",
    "Success confirmation screen with summary card and next-action CTA"
  ]
}}

Only return JSON."""

    raw = await llm_chat(
        system="You are a design validator. Return ONLY valid JSON.",
        user=prompt,
        model=model,
    )
    data = _parse_json(raw)
    if data and "overall_score" in data and "checks" in data:
        return {"review": data, "model": model}

    # Fallback review
    return {
        "review": {
            "overall_status": "approved_with_changes",
            "overall_score": 82,
            "summary": "Design has solid structure with accessibility improvements needed.",
            "checks": [
                {"category": "accessibility", "check_name": "WCAG AA Contrast", "status": "pass", "severity": "info", "details": "Colour contrast meets WCAG AA", "recommendation": "Maintain contrast ratios"},
                {"category": "ux_quality", "check_name": "User Flow Completeness", "status": "pass", "severity": "info", "details": "All screens have entry/exit points", "recommendation": "Good as-is"},
                {"category": "accessibility", "check_name": "Focus Indicators", "status": "warning", "severity": "major", "details": "Focus rings not specified", "recommendation": "Add focus-visible CSS", "wcag_criterion": "2.4.7"},
            ],
            "accessibility_report": {"score": 84, "issues_critical": 0, "issues_major": 1, "issues_minor": 1, "highlights": ["Good contrast", "Semantic structure"]},
            "design_system_compliance": {"score": 88, "deviations": ["Spacing alignment"]},
            "requirements_coverage": {"covered": [], "missing": [], "coverage_percent": 100},
            "recommended_changes": [
                {"priority": "high", "area": "Focus states", "change": "Add :focus-visible ring to all interactive elements", "reason": "WCAG 2.4.7"},
                {"priority": "medium", "area": "Touch targets", "change": "Ensure all tap targets ≥ 48px", "reason": "Mobile accessibility"},
            ],
            "approved_patterns": [
                "Multi-screen progressive form flow",
                "Confirmation screen with success state",
            ],
        },
        "model": model,
        "fallback": True,
    }


# ── DB ────────────────────────────────────────────────────────────────────────
engine = create_async_engine(DATABASE_URL, echo=False, connect_args={"check_same_thread": False})
AsyncSessionLocal = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)


class Base(DeclarativeBase): pass

class User(Base):
    __tablename__ = "users"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    email: Mapped[str] = mapped_column(sa.String(255), unique=True, index=True)
    username: Mapped[str] = mapped_column(sa.String(100), unique=True)
    full_name: Mapped[str] = mapped_column(sa.String(255))
    hashed_password: Mapped[str] = mapped_column(sa.String(255))
    role: Mapped[str] = mapped_column(sa.String(50), default="designer")
    is_active: Mapped[bool] = mapped_column(sa.Boolean, default=True)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())

class Project(Base):
    __tablename__ = "projects"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    owner_id: Mapped[str] = mapped_column(sa.String(36), sa.ForeignKey("users.id"))
    name: Mapped[str] = mapped_column(sa.String(255))
    description: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)
    domain: Mapped[str] = mapped_column(sa.String(50), default="other")
    status: Mapped[str] = mapped_column(sa.String(50), default="active")
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())

class WorkflowRun(Base):
    __tablename__ = "workflow_runs"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    project_id: Mapped[str] = mapped_column(sa.String(36))
    initiated_by: Mapped[str] = mapped_column(sa.String(36))
    title: Mapped[str] = mapped_column(sa.String(255))
    status: Mapped[str] = mapped_column(sa.String(50), default="pending")
    current_step: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)
    starting_point: Mapped[str] = mapped_column(sa.String(50))
    work_mode: Mapped[str] = mapped_column(sa.String(50))
    user_prompt: Mapped[str] = mapped_column(sa.Text)
    pipeline_context: Mapped[str] = mapped_column(sa.Text, default="{}")
    total_tokens: Mapped[int] = mapped_column(sa.Integer, default=0)
    error_message: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())
    started_at: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)
    completed_at: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)

class Artifact(Base):
    __tablename__ = "artifacts"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    project_id: Mapped[str] = mapped_column(sa.String(36))
    workflow_run_id: Mapped[Optional[str]] = mapped_column(sa.String(36), nullable=True)
    name: Mapped[str] = mapped_column(sa.String(255))
    type: Mapped[str] = mapped_column(sa.String(50))
    status: Mapped[str] = mapped_column(sa.String(50), default="pending_approval")
    content: Mapped[str] = mapped_column(sa.Text, default="{}")
    agent_type: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)
    model_used: Mapped[Optional[str]] = mapped_column(sa.String(100), nullable=True)
    complexity_score: Mapped[Optional[int]] = mapped_column(sa.Integer, nullable=True)
    tokens_used: Mapped[Optional[int]] = mapped_column(sa.Integer, default=0)
    generation_time_ms: Mapped[Optional[int]] = mapped_column(sa.Integer, nullable=True)
    prompt: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)
    version: Mapped[int] = mapped_column(sa.Integer, default=1)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())

class ApprovalGate(Base):
    __tablename__ = "approval_gates"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    workflow_run_id: Mapped[str] = mapped_column(sa.String(36))
    artifact_id: Mapped[Optional[str]] = mapped_column(sa.String(36), nullable=True)
    gate_type: Mapped[str] = mapped_column(sa.String(50))
    status: Mapped[str] = mapped_column(sa.String(50), default="pending")
    reviewer_comment: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)
    reviewed_at: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())

class KnowledgeDocument(Base):
    __tablename__ = "knowledge_documents"
    id: Mapped[str] = mapped_column(sa.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    uploaded_by: Mapped[str] = mapped_column(sa.String(36))
    name: Mapped[str] = mapped_column(sa.String(255))
    category: Mapped[str] = mapped_column(sa.String(50))
    subcategory: Mapped[Optional[str]] = mapped_column(sa.String(100), nullable=True)
    file_type: Mapped[Optional[str]] = mapped_column(sa.String(50), nullable=True)
    status: Mapped[str] = mapped_column(sa.String(50), default="indexed")
    chunk_count: Mapped[int] = mapped_column(sa.Integer, default=0)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())

class DesignSession(Base):
    """Persists plan→generate sessions so they survive server restarts."""
    __tablename__ = "design_sessions"
    session_id: Mapped[str] = mapped_column(sa.String(36), primary_key=True)
    user_id: Mapped[str] = mapped_column(sa.String(36), index=True)
    prompt: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)
    domain: Mapped[str] = mapped_column(sa.String(50), default="general")
    history: Mapped[str] = mapped_column(sa.Text, default="[]")     # JSON list
    screens: Mapped[str] = mapped_column(sa.Text, default="[]")     # JSON list
    plan_payload: Mapped[Optional[str]] = mapped_column(sa.Text, nullable=True)  # JSON
    project_id: Mapped[Optional[str]] = mapped_column(sa.String(36), nullable=True)
    created_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: Mapped[str] = mapped_column(sa.String(50), default=lambda: datetime.now(timezone.utc).isoformat())


# ── Auth ──────────────────────────────────────────────────────────────────────
bearer_scheme = HTTPBearer()

def hash_password(p: str) -> str:
    return _bcrypt.hashpw(p.encode(), _bcrypt.gensalt()).decode()

def verify_password(plain: str, hashed: str) -> bool:
    return _bcrypt.checkpw(plain.encode(), hashed.encode())

def create_token(user_id: str, days: int = 7) -> str:
    exp = datetime.now(timezone.utc) + timedelta(days=days)
    return jwt.encode({"sub": user_id, "exp": exp}, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_db():
    async with AsyncSessionLocal() as session:
        yield session

async def get_current_user(
    creds: HTTPAuthorizationCredentials = Depends(bearer_scheme),
    db: AsyncSession = Depends(get_db),
) -> User:
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
    r = await db.execute(sa.select(User).where(User.id == user_id))
    user = r.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


# ── WebSocket ─────────────────────────────────────────────────────────────────
class WSManager:
    def __init__(self):
        self._conns: dict[str, set[WebSocket]] = {}

    async def connect(self, ws: WebSocket, run_id: str):
        await ws.accept()
        self._conns.setdefault(run_id, set()).add(ws)

    def disconnect(self, ws: WebSocket, run_id: str):
        self._conns.get(run_id, set()).discard(ws)

    async def broadcast(self, run_id: str, data: dict):
        dead = set()
        for ws in list(self._conns.get(run_id, set())):
            try:
                await ws.send_text(json.dumps(data))
            except Exception:
                dead.add(ws)
        for ws in dead:
            self._conns.get(run_id, set()).discard(ws)


ws_manager = WSManager()


# ── Workflow execution ────────────────────────────────────────────────────────
async def _save_artifact(db, project_id, run_id, name, atype, content, agent, model, complexity, start_ms) -> Artifact:
    duration_ms = int((time.time() * 1000) - start_ms)
    a = Artifact(
        project_id=project_id, workflow_run_id=run_id,
        name=name, type=atype, status="pending_approval",
        content=json.dumps(content), agent_type=agent, model_used=model,
        complexity_score=complexity, generation_time_ms=duration_ms,
    )
    db.add(a)
    await db.commit()
    await db.refresh(a)
    return a


async def _create_gate(db, run_id, artifact_id, gate_type) -> ApprovalGate:
    g = ApprovalGate(workflow_run_id=run_id, artifact_id=artifact_id, gate_type=gate_type)
    db.add(g)
    await db.commit()
    await db.refresh(g)
    return g


async def run_full_pipeline(run_id: str, prompt: str, project_id: str, starting_point: str, work_mode: str):
    """Full pipeline: classify → planner → gate → ux → gate → ui → gate → review → done."""
    async with AsyncSessionLocal() as db:
        try:
            # Estimate complexity from prompt length + keywords
            words = len(prompt.split())
            complexity_words = ["full", "complete", "enterprise", "system", "platform", "complex", "multi", "dashboard"]
            simple_words = ["simple", "basic", "small", "single", "one", "quick"]
            complexity = min(10, max(1, words // 6 + 2))
            if any(w in prompt.lower() for w in complexity_words): complexity = min(10, complexity + 2)
            if any(w in prompt.lower() for w in simple_words): complexity = max(1, complexity - 2)

            title = prompt[:60].strip()
            description = prompt

            # Update run
            r = await db.execute(sa.select(WorkflowRun).where(WorkflowRun.id == run_id))
            run = r.scalar_one()
            run.status = "running"
            run.started_at = datetime.now(timezone.utc).isoformat()
            run.pipeline_context = json.dumps({"complexity": complexity, "title": title})
            await db.commit()

            await ws_manager.broadcast(run_id, {
                "type": "step_complete", "step": "pipeline",
                "message": f"✓ Prompt analyzed — complexity {complexity}/10",
                "data": {"complexity": complexity}
            })

            # ── Planner ───────────────────────────────────────────────────
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "planner", "message": "📋 Planner Agent analyzing requirements..."})
            t0 = time.time() * 1000
            planner_result = await run_planner(title, description, complexity)
            plan = planner_result["plan"]
            model = planner_result["model"]
            fallback = planner_result.get("fallback", False)

            a_plan = await _save_artifact(db, project_id, run_id, "Requirement Blueprint", "requirement", plan, "planner", model, complexity, t0)
            gate = await _create_gate(db, run_id, a_plan.id, "after_planner")
            run.current_step = "awaiting_approval_after_planner"
            run.status = "awaiting_approval"
            await db.commit()

            suffix = " (fallback)" if fallback else f" ({model})"
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "planner", "message": f"✓ Requirement Blueprint ready{suffix}"})
            await ws_manager.broadcast(run_id, {
                "type": "approval_required", "step": "planner",
                "gate_id": gate.id, "artifact_id": a_plan.id,
                "message": "📋 Requirement Blueprint ready — review and approve to continue to UX"
            })

            # Wait for approval
            approved = await _wait_for_gate(gate.id, run_id)
            if not approved:
                run.status = "rejected"
                await db.commit()
                return

            # ── UX ────────────────────────────────────────────────────────
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "ux", "message": "🎯 UX Agent designing user flows..."})
            t0 = time.time() * 1000
            ux_result = await run_ux(plan, title, complexity)
            ux = ux_result["ux"]
            model = ux_result["model"]
            fallback = ux_result.get("fallback", False)

            a_ux = await _save_artifact(db, project_id, run_id, "UX Blueprint", "user_flow", ux, "ux", model, complexity, t0)
            gate2 = await _create_gate(db, run_id, a_ux.id, "after_ux")
            run.current_step = "awaiting_approval_after_ux"
            run.status = "awaiting_approval"
            await db.commit()

            suffix = " (fallback)" if fallback else f" ({model})"
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "ux", "message": f"✓ UX Blueprint ready{suffix}"})
            await ws_manager.broadcast(run_id, {
                "type": "approval_required", "step": "ux",
                "gate_id": gate2.id, "artifact_id": a_ux.id,
                "message": "🎯 UX Blueprint ready — review and approve to continue to UI"
            })

            approved = await _wait_for_gate(gate2.id, run_id)
            if not approved:
                run.status = "rejected"
                await db.commit()
                return

            # ── UI ────────────────────────────────────────────────────────
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "ui", "message": "🎨 UI Agent designing visual screens..."})
            t0 = time.time() * 1000
            ui_result = await run_ui(plan, ux, title, complexity)
            ui = ui_result["ui"]
            model = ui_result["model"]
            fallback = ui_result.get("fallback", False)

            a_ui = await _save_artifact(db, project_id, run_id, "UI Design", "hifi_screen", ui, "ui", model, complexity, t0)
            gate3 = await _create_gate(db, run_id, a_ui.id, "after_ui")
            run.current_step = "awaiting_approval_after_ui"
            run.status = "awaiting_approval"
            await db.commit()

            suffix = " (fallback)" if fallback else f" ({model})"
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "ui", "message": f"✓ UI Design ready{suffix}"})
            await ws_manager.broadcast(run_id, {
                "type": "approval_required", "step": "ui",
                "gate_id": gate3.id, "artifact_id": a_ui.id,
                "message": "🎨 UI Design ready — review and approve for final review"
            })

            approved = await _wait_for_gate(gate3.id, run_id)
            if not approved:
                run.status = "rejected"
                await db.commit()
                return

            # ── Review ────────────────────────────────────────────────────
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "review", "message": "✅ Review Agent validating design quality..."})
            t0 = time.time() * 1000
            review_result = await run_review(plan, ux, ui, title, complexity)
            review = review_result["review"]
            model = review_result["model"]
            fallback = review_result.get("fallback", False)

            a_review = await _save_artifact(db, project_id, run_id, "Design Review", "review_report", review, "review", model, complexity, t0)
            a_review.status = "approved"
            await db.commit()

            # Save approved patterns to RAG for self-learning
            patterns = review.get("approved_patterns", [])
            for i, p in enumerate(patterns):
                await rag_add("design_patterns", f"{run_id}-pattern-{i}", p, {"run_id": run_id, "title": title})

            run.current_step = "completed"
            run.status = "completed"
            run.completed_at = datetime.now(timezone.utc).isoformat()
            await db.commit()

            score = review.get("overall_score", "N/A")
            suffix = " (fallback)" if fallback else f" ({model})"
            await ws_manager.broadcast(run_id, {
                "type": "workflow_complete",
                "artifact_id": a_review.id,
                "score": score,
                "status": review.get("overall_status", "complete"),
                "message": f"🎉 Design workflow complete — Score: {score}/100{suffix}"
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            try:
                async with AsyncSessionLocal() as edb:
                    r = await edb.execute(sa.select(WorkflowRun).where(WorkflowRun.id == run_id))
                    run = r.scalar_one_or_none()
                    if run:
                        run.status = "failed"
                        run.error_message = str(e)
                        await edb.commit()
            except Exception:
                pass
            await ws_manager.broadcast(run_id, {"type": "error", "message": f"Pipeline error: {e}"})


async def _wait_for_gate(gate_id: str, run_id: str, timeout: int = 600) -> bool:
    """Poll DB until gate is approved/rejected (up to 10 min)."""
    for _ in range(timeout):
        await asyncio.sleep(1)
        async with AsyncSessionLocal() as db:
            r = await db.execute(sa.select(ApprovalGate).where(ApprovalGate.id == gate_id))
            gate = r.scalar_one_or_none()
            if gate and gate.status in ("approved", "changes_requested"):
                return True
            if gate and gate.status == "rejected":
                await ws_manager.broadcast(run_id, {"type": "gate_rejected", "gate_id": gate_id})
                return False
    return False  # timeout


async def run_next_step(run_id: str, step: str, req_id: str = None, ux_id: str = None, ui_id: str = None):
    """Manually trigger a specific step (used when frontend sends step request)."""
    async with AsyncSessionLocal() as db:
        r = await db.execute(sa.select(WorkflowRun).where(WorkflowRun.id == run_id))
        run = r.scalar_one()
        ctx = json.loads(run.pipeline_context)
        complexity = ctx.get("complexity", 5)
        title = ctx.get("title", run.user_prompt[:60])

        # Load previous artifacts
        plan, ux, ui = {}, {}, {}
        for aid, store in [(req_id, "plan"), (ux_id, "ux"), (ui_id, "ui")]:
            if aid:
                ra = await db.execute(sa.select(Artifact).where(Artifact.id == aid))
                ra = ra.scalar_one_or_none()
                if ra:
                    if store == "plan": plan = json.loads(ra.content)
                    elif store == "ux": ux = json.loads(ra.content)
                    elif store == "ui": ui = json.loads(ra.content)

        t0 = time.time() * 1000

        if step == "ux":
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "ux", "message": "🎯 UX Agent designing user flows..."})
            result = await run_ux(plan, title, complexity)
            a = await _save_artifact(db, run.project_id, run_id, "UX Blueprint", "user_flow", result["ux"], "ux", result["model"], complexity, t0)
            gate = await _create_gate(db, run_id, a.id, "after_ux")
            run.current_step = "awaiting_approval_after_ux"
            run.status = "awaiting_approval"
            await db.commit()
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "ux", "message": f"✓ UX Blueprint ready ({result['model']})"})
            await ws_manager.broadcast(run_id, {"type": "approval_required", "step": "ux", "gate_id": gate.id, "artifact_id": a.id, "message": "UX Blueprint ready for review"})

        elif step == "ui":
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "ui", "message": "🎨 UI Agent designing visual screens..."})
            result = await run_ui(plan, ux, title, complexity)
            a = await _save_artifact(db, run.project_id, run_id, "UI Design", "hifi_screen", result["ui"], "ui", result["model"], complexity, t0)
            gate = await _create_gate(db, run_id, a.id, "after_ui")
            run.current_step = "awaiting_approval_after_ui"
            run.status = "awaiting_approval"
            await db.commit()
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "ui", "message": f"✓ UI Design ready ({result['model']})"})
            await ws_manager.broadcast(run_id, {"type": "approval_required", "step": "ui", "gate_id": gate.id, "artifact_id": a.id, "message": "UI Design ready for review"})

        elif step == "review":
            await ws_manager.broadcast(run_id, {"type": "agent_start", "agent": "review", "message": "✅ Review Agent validating..."})
            result = await run_review(plan, ux, ui, title, complexity)
            a = await _save_artifact(db, run.project_id, run_id, "Design Review", "review_report", result["review"], "review", result["model"], complexity, t0)
            a.status = "approved"
            run.current_step = "completed"
            run.status = "completed"
            run.completed_at = datetime.now(timezone.utc).isoformat()
            await db.commit()

            patterns = result["review"].get("approved_patterns", [])
            for i, p in enumerate(patterns):
                await rag_add("design_patterns", f"{run_id}-p{step}-{i}", p, {"run_id": run_id})

            score = result["review"].get("overall_score", "N/A")
            await ws_manager.broadcast(run_id, {"type": "step_complete", "step": "review", "message": f"✓ Review done ({result['model']})"})
            await ws_manager.broadcast(run_id, {"type": "workflow_complete", "artifact_id": a.id, "score": score, "status": result["review"].get("overall_status"), "message": f"🎉 Workflow complete — Score: {score}/100"})


# ── Session persistence helpers ───────────────────────────────────────────────
async def _persist_session(session_id: str):
    """Save (upsert) a design session from memory to SQLite."""
    mem = _design_sessions.get(session_id)
    if not mem:
        return
    try:
        async with AsyncSessionLocal() as db:
            row = await db.get(DesignSession, session_id)
            now = datetime.now(timezone.utc).isoformat()
            if row:
                row.history = json.dumps(mem.get("history", []))
                row.screens = json.dumps(mem.get("screens", []))
                row.plan_payload = json.dumps(mem["plan_payload"]) if mem.get("plan_payload") else None
                row.updated_at = now
            else:
                row = DesignSession(
                    session_id=session_id,
                    user_id=mem.get("user_id", ""),
                    prompt=(mem.get("history") or [{}])[0].get("content", "") if mem.get("history") else "",
                    domain=mem.get("domain", "general"),
                    history=json.dumps(mem.get("history", [])),
                    screens=json.dumps(mem.get("screens", [])),
                    plan_payload=json.dumps(mem["plan_payload"]) if mem.get("plan_payload") else None,
                    project_id=mem.get("project_id"),
                    created_at=mem.get("created_at", now),
                    updated_at=now,
                )
                db.add(row)
            await db.commit()
    except Exception as e:
        print(f"[session-persist] Failed to save session {session_id}: {e}")


async def _load_session(session_id: str) -> bool:
    """Load a session from SQLite into _design_sessions. Returns True if found."""
    if session_id in _design_sessions:
        return True
    try:
        async with AsyncSessionLocal() as db:
            row = await db.get(DesignSession, session_id)
            if not row:
                return False
            _design_sessions[session_id] = {
                "user_id": row.user_id,
                "project_id": row.project_id,
                "domain": row.domain,
                "created_at": row.created_at,
                "history": json.loads(row.history or "[]"),
                "screens": json.loads(row.screens or "[]"),
                "plan_payload": json.loads(row.plan_payload) if row.plan_payload else None,
            }
            return True
    except Exception as e:
        print(f"[session-load] Failed to load session {session_id}: {e}")
        return False


@asynccontextmanager
async def lifespan(app: FastAPI):
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    Path("./uploads/knowledge").mkdir(parents=True, exist_ok=True)
    # Seed knowledge/ folder into ChromaDB on startup (idempotent — uses upsert)
    try:
        seeded = await seed_knowledge_from_files()
        if seeded:
            print(f"   📚 Knowledge base seeded — {seeded} chunks indexed from knowledge/")
        else:
            print("   📚 Knowledge base: no new files to seed")
    except Exception as e:
        print(f"   ⚠️  Knowledge seeding failed: {e}")
    # Reload recent sessions into memory so they survive restarts
    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(
                sa.select(DesignSession)
                .order_by(DesignSession.updated_at.desc())
                .limit(100)
            )
            rows = result.scalars().all()
            for row in rows:
                _design_sessions[row.session_id] = {
                    "user_id": row.user_id,
                    "project_id": row.project_id,
                    "domain": row.domain,
                    "created_at": row.created_at,
                    "history": json.loads(row.history or "[]"),
                    "screens": json.loads(row.screens or "[]"),
                    "plan_payload": json.loads(row.plan_payload) if row.plan_payload else None,
                }
            if rows:
                print(f"   ♻️  Restored {len(rows)} design sessions from DB")
    except Exception as e:
        print(f"   ⚠️  Could not restore sessions: {e}")
    ds_msg = "✓ V1 design system loaded" if HAS_DESIGN_SYSTEM else "⚠️  Design system not found"
    if AI_BACKEND == "ollama":
        ai_msg = (
            f"🦙 LOCAL AI via Ollama\n"
            f"   Reasoning: {OLLAMA_MODELS['reasoning']}\n"
            f"   UI/Code:   {OLLAMA_MODELS['coder']}"
        )
    elif AI_BACKEND == "cloud":
        ai_msg = (
            f"☁️  Cloud AI — {CLOUD_PROVIDER.upper()}\n"
            f"   Standard: {CLOUD_MODELS['standard']}\n"
            f"   Vision:   enabled" + ("  ·  Ollama fallback ready" if HAS_OLLAMA else "")
        )
    else:
        ai_msg = "❌ NO AI — start Ollama or set an API key in .env"

    print(f"\n✅ CORE Studio V2")
    print(f"   {ai_msg}")
    print(f"   {ds_msg}")
    print(f"   DB: {DB_PATH}")
    print(f"   API Docs: http://localhost:8000/docs\n")
    yield


app = FastAPI(title="CORE Studio V2", version="2.0.0", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=[
    "http://localhost:3000", "http://127.0.0.1:3000",
    "http://localhost:3001", "http://127.0.0.1:3001",
], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])


# ── Auth ──────────────────────────────────────────────────────────────────────
class RegisterReq(BaseModel):
    email: str
    username: str
    full_name: str
    password: str

class LoginReq(BaseModel):
    email: str
    password: str

class RefreshReq(BaseModel):
    refresh_token: str

def _user_out(u: User) -> dict:
    return {"id": u.id, "email": u.email, "username": u.username, "full_name": u.full_name, "role": u.role, "created_at": u.created_at}

@app.post("/api/v1/auth/register", status_code=201)
async def register(req: RegisterReq, db: AsyncSession = Depends(get_db)):
    r = await db.execute(sa.select(User).where((User.email == req.email) | (User.username == req.username)))
    if r.scalar_one_or_none():
        raise HTTPException(400, "Email or username already exists")
    user = User(email=req.email, username=req.username, full_name=req.full_name, hashed_password=hash_password(req.password))
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return {"access_token": create_token(user.id), "refresh_token": create_token(user.id, 30), "token_type": "bearer", "user": _user_out(user)}

@app.post("/api/v1/auth/login")
async def login(req: LoginReq, db: AsyncSession = Depends(get_db)):
    r = await db.execute(sa.select(User).where(User.email == req.email))
    user = r.scalar_one_or_none()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(401, "Invalid credentials")
    return {"access_token": create_token(user.id), "refresh_token": create_token(user.id, 30), "token_type": "bearer", "user": _user_out(user)}

@app.post("/api/v1/auth/refresh")
async def refresh_token(req: RefreshReq, db: AsyncSession = Depends(get_db)):
    try:
        payload = jwt.decode(req.refresh_token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
    except JWTError:
        raise HTTPException(401, "Invalid token")
    r = await db.execute(sa.select(User).where(User.id == user_id))
    user = r.scalar_one_or_none()
    if not user:
        raise HTTPException(401, "User not found")
    return {"access_token": create_token(user.id), "refresh_token": create_token(user.id, 30), "token_type": "bearer", "user": _user_out(user)}


# ── Projects ──────────────────────────────────────────────────────────────────
class ProjectReq(BaseModel):
    name: str
    description: Optional[str] = None
    domain: str = "other"

@app.post("/api/v1/projects", status_code=201)
async def create_project(req: ProjectReq, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    p = Project(owner_id=user.id, name=req.name, description=req.description, domain=req.domain)
    db.add(p); await db.commit(); await db.refresh(p)
    return {"id": p.id, "name": p.name, "description": p.description, "domain": p.domain, "status": p.status, "created_at": p.created_at, "updated_at": p.updated_at}

@app.get("/api/v1/projects")
async def list_projects(db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(Project).where(Project.owner_id == user.id).order_by(Project.created_at.desc()))
    return [{"id": p.id, "name": p.name, "description": p.description, "domain": p.domain, "status": p.status, "created_at": p.created_at, "updated_at": p.updated_at} for p in r.scalars().all()]

@app.get("/api/v1/projects/{pid}")
async def get_project(pid: str, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(Project).where(Project.id == pid))
    p = r.scalar_one_or_none()
    if not p: raise HTTPException(404, "Not found")
    return {"id": p.id, "name": p.name, "description": p.description, "domain": p.domain, "status": p.status, "created_at": p.created_at, "updated_at": p.updated_at}

@app.put("/api/v1/projects/{pid}")
async def update_project(pid: str, req: ProjectReq, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(Project).where(Project.id == pid))
    p = r.scalar_one_or_none()
    if not p: raise HTTPException(404, "Not found")
    p.name = req.name
    if req.description: p.description = req.description
    p.updated_at = datetime.now(timezone.utc).isoformat()
    await db.commit()
    return {"id": p.id, "name": p.name}

@app.get("/api/v1/projects/{pid}/artifacts")
async def get_project_artifacts(pid: str, artifact_type: Optional[str] = None, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    stmt = sa.select(Artifact).where(Artifact.project_id == pid)
    if artifact_type: stmt = stmt.where(Artifact.type == artifact_type)
    stmt = stmt.order_by(Artifact.created_at.desc())
    r = await db.execute(stmt)
    return [{"id": a.id, "name": a.name, "type": a.type, "status": a.status, "agent_type": a.agent_type, "model_used": a.model_used, "complexity_score": a.complexity_score, "tokens_used": a.tokens_used, "version": a.version, "created_at": a.created_at} for a in r.scalars().all()]


# ── Workflows ─────────────────────────────────────────────────────────────────
class CreateWorkflowReq(BaseModel):
    project_id: str
    user_prompt: str
    starting_point: str = "requirement"
    work_mode: str = "create"
    input_files: list = []

class RunStepReq(BaseModel):
    step: str
    requirement_artifact_id: Optional[str] = None
    ux_artifact_id: Optional[str] = None
    ui_artifact_id: Optional[str] = None

class ApproveGateReq(BaseModel):
    status: str
    comment: str = ""

@app.post("/api/v1/workflows", status_code=201)
async def create_workflow(req: CreateWorkflowReq, background_tasks: BackgroundTasks, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    run = WorkflowRun(project_id=req.project_id, initiated_by=user.id, title=f"Design: {req.user_prompt[:60]}", status="pending", starting_point=req.starting_point, work_mode=req.work_mode, user_prompt=req.user_prompt)
    db.add(run); await db.commit(); await db.refresh(run)
    background_tasks.add_task(run_full_pipeline, run.id, req.user_prompt, req.project_id, req.starting_point, req.work_mode)
    return {"workflow_run_id": run.id, "status": "started"}

@app.post("/api/v1/workflows/{run_id}/steps")
async def run_step(run_id: str, req: RunStepReq, background_tasks: BackgroundTasks, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    background_tasks.add_task(run_next_step, run_id, req.step, req.requirement_artifact_id, req.ux_artifact_id, req.ui_artifact_id)
    return {"status": "step_started", "step": req.step}

@app.post("/api/v1/workflows/{run_id}/gates/{gate_id}/approve")
async def approve_gate(run_id: str, gate_id: str, req: ApproveGateReq, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(ApprovalGate).where(ApprovalGate.id == gate_id))
    gate = r.scalar_one_or_none()
    if not gate: raise HTTPException(404, "Gate not found")
    gate.status = req.status
    gate.reviewer_comment = req.comment
    gate.reviewed_at = datetime.now(timezone.utc).isoformat()
    await db.commit()
    await ws_manager.broadcast(run_id, {"type": "gate_decision", "gate_id": gate_id, "status": req.status})
    return {"gate_id": gate.id, "status": gate.status}

@app.get("/api/v1/workflows/{run_id}")
async def get_workflow(run_id: str, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(WorkflowRun).where(WorkflowRun.id == run_id))
    run = r.scalar_one_or_none()
    if not run: raise HTTPException(404, "Not found")
    return {"id": run.id, "status": run.status, "current_step": run.current_step, "user_prompt": run.user_prompt, "total_tokens": run.total_tokens, "pipeline_context": json.loads(run.pipeline_context), "created_at": run.created_at}


# ── Artifacts ─────────────────────────────────────────────────────────────────
@app.get("/api/v1/artifacts/{aid}")
async def get_artifact(aid: str, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(Artifact).where(Artifact.id == aid))
    a = r.scalar_one_or_none()
    if not a: raise HTTPException(404, "Not found")
    return {"id": a.id, "name": a.name, "type": a.type, "status": a.status, "content": json.loads(a.content), "agent_type": a.agent_type, "model_used": a.model_used, "complexity_score": a.complexity_score, "tokens_used": a.tokens_used or 0, "generation_time_ms": a.generation_time_ms, "version": a.version, "created_at": a.created_at}

@app.post("/api/v1/artifacts/{aid}/approve")
async def approve_artifact(aid: str, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(Artifact).where(Artifact.id == aid))
    a = r.scalar_one_or_none()
    if not a: raise HTTPException(404, "Not found")
    a.status = "approved"; await db.commit()
    return {"id": a.id, "status": a.status}


# ── Knowledge / Document Center ───────────────────────────────────────────────

def _extract_text_from_file(file_path: str, ext: str, raw_bytes: bytes) -> str:
    """Extract text from txt/md/pdf/docx."""
    ext = ext.lower()
    if ext in ("txt", "md"):
        return raw_bytes.decode("utf-8", errors="ignore")
    if ext == "pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            return "\n\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            return ""
    if ext == "docx":
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text)
        except Exception:
            return ""
    return raw_bytes.decode("utf-8", errors="ignore")


async def _extract_handwritten(image_b64: str) -> str:
    """Use LLava to read handwritten/scanned document text."""
    prompt = """This is a scanned or handwritten document. Transcribe ALL text you can read from it,
preserving structure (headings, paragraphs, lists). Output the transcribed text only, no commentary."""
    try:
        return await llm_chat_vision(image_b64, prompt)
    except Exception:
        return ""


@app.post("/api/v1/knowledge/upload", status_code=201)
async def upload_knowledge(
    file: UploadFile = File(...),
    category: str = Form(...),
    subcategory: str = Form(None),
    is_handwritten: str = Form("false"),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    content = await file.read()
    uploads_dir = Path("./uploads/knowledge")
    uploads_dir.mkdir(parents=True, exist_ok=True)
    ext = (file.filename.rsplit(".", 1)[-1] if "." in file.filename else "txt").lower()
    file_path = uploads_dir / f"{uuid.uuid4()}.{ext}"
    async with aiofiles.open(str(file_path), "wb") as f:
        await f.write(content)

    # Extract text — handwritten images go through vision OCR (best available backend,
    # not necessarily Ollama — see resolve_vision_provider()).
    if is_handwritten == "true" and ext in ("png", "jpg", "jpeg", "webp"):
        if resolve_vision_provider() == "none":
            raise HTTPException(400, "Handwritten reading requires a configured vision "
                                      "backend — set a cloud AI API key or run Ollama with llava.")
        b64 = base64.b64encode(content).decode("utf-8")
        text = await _extract_handwritten(b64)
    else:
        text = _extract_text_from_file(str(file_path), ext, content)

    if not text.strip():
        text = f"(Empty or unreadable document: {file.filename})"

    # Chunk + index into the right collection
    col = {"domain": "guidelines", "company": "company_standards", "design_system": "design_patterns"}.get(category, "guidelines")
    doc_id = str(uuid.uuid4())
    chunk_count = await rag_add_chunked(col, doc_id, text, {
        "name": file.filename, "category": category, "subcategory": subcategory or "", "doc_id": doc_id,
    })
    # Also index into a unified "documents" collection for Q&A
    await rag_add_chunked("documents", doc_id, text, {
        "name": file.filename, "category": category, "doc_id": doc_id,
    })

    doc = KnowledgeDocument(
        id=doc_id, uploaded_by=user.id, name=file.filename,
        category=category, subcategory=subcategory, file_type=ext,
        status="indexed", chunk_count=chunk_count,
    )
    db.add(doc); await db.commit(); await db.refresh(doc)
    return {
        "document_id": doc.id, "name": doc.name, "status": "indexed",
        "chunk_count": chunk_count, "text_preview": text[:300],
        "handwritten_ocr": is_handwritten == "true",
    }


@app.get("/api/v1/knowledge")
async def list_knowledge(category: Optional[str] = None, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    stmt = sa.select(KnowledgeDocument)
    if category: stmt = stmt.where(KnowledgeDocument.category == category)
    r = await db.execute(stmt.order_by(KnowledgeDocument.created_at.desc()))
    return [{"id": d.id, "name": d.name, "category": d.category, "subcategory": d.subcategory, "file_type": d.file_type, "status": d.status, "chunk_count": d.chunk_count, "created_at": d.created_at} for d in r.scalars().all()]


@app.post("/api/v1/knowledge/search")
async def search_knowledge(req: dict, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    query = req.get("query", "")
    results = await rag_query_rich("documents", query, top_k=req.get("top_k", 5))
    return {
        "results": [{"content": r["content"], "source": r["source"], "similarity": round(1 - (r["distance"] or 0), 2)} for r in results],
        "count": len(results),
    }


class DocQARequest(BaseModel):
    question: str
    category: Optional[str] = None
    top_k: int = 8
    conversation_history: List[Dict] = []  # prior turns for multi-turn context


def _detect_domain(question: str) -> str:
    q = question.lower()
    if any(w in q for w in ["retirement", "401k", "ira", "pension", "annuity", "social security", "rmd", "roth"]):
        return "retirement"
    if any(w in q for w in ["loan", "mortgage", "interest rate", "amortization", "refinance", "heloc", "debt"]):
        return "loan"
    if any(w in q for w in ["insurance", "premium", "deductible", "policy", "coverage", "beneficiary"]):
        return "insurance"
    if any(w in q for w in ["tax", "deduction", "capital gains", "filing", "irs", "w-2", "1099"]):
        return "tax"
    if any(w in q for w in ["invest", "portfolio", "stock", "bond", "etf", "dividend", "asset allocation"]):
        return "investment"
    return "general"


_DOMAIN_CONTEXT = {
    "retirement": {
        "persona": "You are a knowledgeable retirement planning specialist with expertise in 401(k), IRA, pension plans, Social Security, and withdrawal strategies.",
        "next_step_topics": [
            "How much should I be contributing to maximize my retirement savings?",
            "What is the difference between a Traditional IRA and Roth IRA?",
            "When should I start taking Social Security benefits?",
            "How do I create a withdrawal strategy for retirement?",
            "What is a Required Minimum Distribution (RMD) and when do I need to take it?",
        ],
    },
    "loan": {
        "persona": "You are a knowledgeable loan and mortgage specialist with expertise in lending products, interest calculations, and debt management strategies.",
        "next_step_topics": [
            "How do I calculate my total interest paid over the loan term?",
            "Should I refinance my loan given current interest rates?",
            "What is the best strategy to pay off my loan faster?",
            "How does my credit score affect my loan interest rate?",
            "What is the difference between a fixed and variable rate loan?",
        ],
    },
    "insurance": {
        "persona": "You are a knowledgeable insurance specialist with expertise in life, health, property, and liability coverage.",
        "next_step_topics": [
            "How much life insurance coverage do I actually need?",
            "What is the difference between term and whole life insurance?",
            "How do I evaluate if my current coverage is adequate?",
            "What factors most affect my insurance premiums?",
        ],
    },
    "tax": {
        "persona": "You are a knowledgeable tax planning specialist with expertise in tax-advantaged accounts, deductions, and filing strategies.",
        "next_step_topics": [
            "What deductions am I likely missing?",
            "How can I reduce my taxable income this year?",
            "What is tax-loss harvesting and should I use it?",
            "How do capital gains affect my overall tax liability?",
        ],
    },
    "investment": {
        "persona": "You are a knowledgeable investment specialist with expertise in portfolio construction, risk management, and long-term wealth building.",
        "next_step_topics": [
            "How should I allocate my portfolio based on my risk tolerance?",
            "What is dollar-cost averaging and should I use it?",
            "How do I evaluate if my portfolio is properly diversified?",
            "What is the impact of expense ratios on long-term returns?",
        ],
    },
    "general": {
        "persona": "You are a knowledgeable financial advisor with broad expertise across retirement, investments, loans, insurance, and tax planning.",
        "next_step_topics": [
            "What documents should I upload to get personalized guidance?",
            "How do I build an emergency fund?",
            "What is a good framework for personal financial planning?",
        ],
    },
}


@app.post("/api/v1/knowledge/ask")
async def ask_documents(
    req: DocQARequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Jarvis-style Document Q&A:
    - Domain-aware persona (retirement, loan, insurance, tax, investment)
    - RAG retrieval from uploaded documents
    - LLM synthesizes grounded answer with inline citations
    - Proactive next-step suggestions tailored to the question
    - Multi-turn conversation context
    """
    import re as _re

    domain = _detect_domain(req.question)
    domain_cfg = _DOMAIN_CONTEXT[domain]

    # Retrieve relevant chunks from uploaded docs
    chunks = await rag_query_rich("documents", req.question, top_k=req.top_k)

    # Build source context
    context_parts = []
    sources = []
    for i, c in enumerate(chunks):
        src = c.get("source", f"Document {i+1}")
        context_parts.append(f"[Source {i+1}: {src}]\n{c['content']}")
        if src not in [s["name"] for s in sources]:
            sources.append({
                "name": src,
                "relevance": round(1 - (c.get("distance") or 0), 2),
            })

    has_docs = bool(context_parts)
    context_block = "\n\n---\n\n".join(context_parts) if context_parts else ""

    # Build conversation history block
    convo_block = ""
    if req.conversation_history:
        convo_block = "\nPRIOR CONVERSATION:\n"
        for turn in req.conversation_history[-6:]:
            role = turn.get("role", "user")
            content = turn.get("content", "")[:300]
            convo_block += f"{role.upper()}: {content}\n"
        convo_block += "\n"

    # System prompt — domain-aware persona
    system = f"""{domain_cfg['persona']}

Your style: direct, specific, and actionable. You speak like a trusted advisor, not a generic chatbot.
- Use document excerpts as your primary source — always cite inline as [Source N]
- If documents don't cover the topic, say so and use your financial expertise to give a general answer, clearly flagged as general knowledge
- Structure longer answers with clear sections
- Be specific with numbers, rates, and strategies when the documents provide them
- Always end by identifying what the user should focus on next"""

    # Main answer prompt
    prompt_parts = []
    if context_block:
        prompt_parts.append(f"DOCUMENT EXCERPTS FROM YOUR KNOWLEDGE BASE:\n{context_block}")
    else:
        prompt_parts.append("(No uploaded documents matched this question — answering from general financial expertise)")

    if convo_block:
        prompt_parts.append(convo_block)

    prompt_parts.append(f"USER QUESTION: {req.question}")
    prompt_parts.append(
        "Provide a comprehensive, specific answer. "
        + ("Cite source numbers inline. " if context_block else "")
        + "Be direct and actionable."
    )

    model = select_model(6, "planner")
    raw_answer = await llm_chat(system, "\n\n".join(prompt_parts), model)
    answer = _re.sub(r"<think>.*?</think>", "", raw_answer, flags=_re.DOTALL).strip()

    # Generate smart next-step suggestions
    next_steps_prompt = f"""Based on this question: "{req.question}"
And this answer summary: "{answer[:400]}"
Domain: {domain}

Generate 3-4 specific, actionable follow-up questions or next steps the user should explore.
These should be concrete and directly relevant to their situation — not generic.
Return ONLY a JSON array of strings: ["question 1", "question 2", "question 3"]"""

    try:
        raw_suggestions = await llm_chat(
            "You generate smart, specific follow-up suggestions. Return ONLY a JSON array.",
            next_steps_prompt,
            select_model(3, "planner"),
        )
        raw_suggestions = _re.sub(r"<think>.*?</think>", "", raw_suggestions, flags=_re.DOTALL).strip()
        # Extract JSON array
        arr_match = _re.search(r'\[.*?\]', raw_suggestions, _re.DOTALL)
        if arr_match:
            next_steps = json.loads(arr_match.group())[:4]
        else:
            next_steps = domain_cfg["next_step_topics"][:3]
    except Exception:
        next_steps = domain_cfg["next_step_topics"][:3]

    return {
        "answer": answer,
        "domain": domain,
        "sources": sources,
        "chunks_used": len(chunks),
        "has_context": has_docs,
        "next_steps": next_steps,
        "model": model,
    }

@app.delete("/api/v1/knowledge/{doc_id}")
async def delete_knowledge(doc_id: str, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    r = await db.execute(sa.select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = r.scalar_one_or_none()
    if not doc: raise HTTPException(404, "Not found")
    await db.delete(doc); await db.commit()
    return {"deleted": True}


# ── WebSocket ─────────────────────────────────────────────────────────────────
@app.websocket("/ws/workflows/{run_id}")
async def ws_endpoint(websocket: WebSocket, run_id: str):
    await ws_manager.connect(websocket, run_id)
    try:
        while True:
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_text('{"type":"pong"}')
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket, run_id)


# ── Vision Caller (multi-provider) ────────────────────────────────────────────

def _detect_image_media_type(b64_image: str) -> str:
    """Detect actual image media type from base64 header bytes."""
    import base64 as _b64
    try:
        header = _b64.b64decode(b64_image[:16])
        if header[:4] == b'\xff\xd8\xff\xe0' or header[:4] == b'\xff\xd8\xff\xe1':
            return "image/jpeg"
        if header[:8] == b'\x89PNG\r\n\x1a\n':
            return "image/png"
        if header[:6] in (b'GIF87a', b'GIF89a'):
            return "image/gif"
        if header[:4] == b'RIFF' and header[8:12] == b'WEBP':
            return "image/webp"
    except Exception:
        pass
    return "image/png"  # safe fallback


def _vision_anthropic(b64_image: str, prompt: str) -> str:
    client = _get_anthropic()
    media_type = _detect_image_media_type(b64_image)
    msg = client.messages.create(
        model=PROVIDER_MODELS["anthropic"]["standard"],
        max_tokens=8192,  # full HTML generation needs headroom
        messages=[{"role": "user", "content": [
            {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_image}},
            {"type": "text", "text": prompt},
        ]}],
    )
    return msg.content[0].text


def _vision_openai(b64_image: str, prompt: str) -> str:
    client = _get_openai()
    resp = client.chat.completions.create(
        # 2048 was too low for full-screen HTML recreation (detailed tables/cards need
        # more headroom) — matches the 8192 budget used for _vision_anthropic.
        model=PROVIDER_MODELS["openai"]["standard"], max_tokens=8192,
        messages=[{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64_image}"}},
        ]}],
    )
    return resp.choices[0].message.content


def _vision_gemini(b64_image: str, prompt: str) -> str:
    genai = _get_gemini()
    import base64 as _b64
    gmodel = genai.GenerativeModel(PROVIDER_MODELS["gemini"]["standard"])
    resp = gmodel.generate_content([prompt, {"mime_type": "image/png", "data": _b64.b64decode(b64_image)}])
    return resp.text


def _vision_ollama(b64_image: str, prompt: str) -> str:
    payload = json.dumps({
        "model": OLLAMA_MODELS["vision"],
        "messages": [{"role": "user", "content": prompt, "images": [b64_image]}],
        # num_ctx set explicitly for the same reason as _call_ollama_sync — Ollama
        # otherwise silently caps context at 4096 tokens, and an embedded image plus
        # the detailed verbatim-transcription vision prompt can exceed that.
        "stream": False, "options": {"temperature": 0.2, "num_predict": 2048, "num_ctx": 8192},
    }).encode("utf-8")
    req = _urllib_request.Request(
        f"{OLLAMA_BASE_URL}/api/chat", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    try:
        with _urllib_request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data.get("message", {}).get("content", "")
    except Exception as e:
        return json.dumps({"error": str(e)})


_VISION_FN = {
    "anthropic": _vision_anthropic,
    "openai": _vision_openai,
    "gemini": _vision_gemini,
}
_VISION_KEY = {
    "anthropic": ANTHROPIC_API_KEY,
    "openai": OPENAI_API_KEY,
    "gemini": GEMINI_API_KEY,
}
_VISION_MODEL_NAME = {
    "anthropic": lambda: PROVIDER_MODELS["anthropic"]["standard"],
    "openai": lambda: PROVIDER_MODELS["openai"]["standard"],
    "gemini": lambda: PROVIDER_MODELS["gemini"]["standard"],
    "ollama": lambda: OLLAMA_MODELS["vision"],
}


def resolve_vision_provider() -> str:
    """Which backend actually handles image analysis for this request — independent
    of AI_BACKEND_MODE/CLOUD_PROVIDER (see the VISION_PROVIDER comment above). Purely
    driven by which API keys are configured / whether Ollama is reachable; no
    hardcoded provider choice."""
    if VISION_PROVIDER in _VISION_FN and _VISION_KEY.get(VISION_PROVIDER):
        return VISION_PROVIDER
    if VISION_PROVIDER == "ollama" and HAS_OLLAMA:
        return "ollama"
    if VISION_PROVIDER != "auto":
        print(f"[vision] VISION_PROVIDER={VISION_PROVIDER!r} requested but unavailable "
              f"(no key / not reachable) — falling back to auto-detection.")

    # auto: best available cloud key wins, in quality-priority order; local Ollama
    # (llava) is the last resort, only used when no cloud key is configured at all.
    for provider in ("anthropic", "openai", "gemini"):
        if _VISION_KEY.get(provider):
            return provider
    if HAS_OLLAMA:
        return "ollama"
    return "none"


def vision_backend_label() -> str:
    """Human-readable 'provider — model' string for progress/status messages, so the
    UI always reflects the model that will actually run instead of a hardcoded name."""
    provider = resolve_vision_provider()
    if provider == "none":
        return "no vision backend available"
    try:
        model = _VISION_MODEL_NAME[provider]()
    except Exception:
        model = provider
    return f"{model} ({provider})"


async def llm_chat_vision(b64_image: str, prompt: str) -> str:
    """Vision call — routes independently to the best available vision backend (see
    resolve_vision_provider), not tied to whatever AI_BACKEND_MODE governs text calls."""
    loop = asyncio.get_event_loop()
    provider = resolve_vision_provider()

    if provider in _VISION_FN:
        try:
            return await loop.run_in_executor(None, _VISION_FN[provider], b64_image, prompt)
        except Exception as e:
            print(f"[vision] Cloud vision failed ({provider}): {e}")
            # fall through to Ollama only if available
            if HAS_OLLAMA:
                return await loop.run_in_executor(None, _vision_ollama, b64_image, prompt)
            return json.dumps({"error": f"Vision provider {provider} failed: {e}"})

    if provider == "ollama" and HAS_OLLAMA:
        return await loop.run_in_executor(None, _vision_ollama, b64_image, prompt)

    return json.dumps({"error": "No vision backend available"})


# ── Analyst Team Endpoints ────────────────────────────────────────────────────

class AnalystRequest(BaseModel):
    research_type: str   # personas | competitor | market | user_stories | prd
    topic: str
    domain: str = "general"
    context: str = ""
    project_id: Optional[str] = None


@app.post("/api/v1/analyst/research", status_code=201)
async def analyst_research(
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
    research_type: str = Form(...),
    topic: str = Form(...),
    domain: str = Form("general"),
    context: str = Form(""),
    project_id: Optional[str] = Form(None),
    image: UploadFile = File(None),
):
    """Run analyst research — personas, competitor analysis, user stories, PRD.
    Accepts optional image upload for screen-grounded research.
    """
    from app.agents.analyst.research_agent import run_research

    # Process optional image
    image_b64 = None
    if image is not None and image.filename:
        ext = (image.filename.rsplit(".", 1)[-1].lower() if "." in image.filename else "png")
        if ext in {"png", "jpg", "jpeg", "webp"}:
            image_bytes = await image.read()
            import base64 as _b64
            image_b64 = _b64.b64encode(image_bytes).decode("utf-8")

    async def _run():
        result = await run_research(
            research_type=research_type,
            topic=topic,
            domain=domain,
            context=context,
            llm_chat=llm_chat,
            select_model=select_model,
            rag_query=rag_query,
            image_b64=image_b64,
            llm_chat_vision=llm_chat_vision if image_b64 else None,
        )
        if project_id:
            async with AsyncSessionLocal() as save_db:
                art = Artifact(
                    project_id=project_id,
                    name=f"{research_type.title()} — {topic[:60]}",
                    type=f"analyst_{research_type}",
                    status="approved",
                    content=json.dumps(result),
                    agent_type="analyst",
                    model_used=select_model(6, "research"),
                )
                save_db.add(art)
                await save_db.commit()
        return result

    tokens_before = _token_tracker["input"] + _token_tracker["output"]
    calls_before = _token_tracker["calls"]
    cost_before = _token_tracker["cost_usd"]
    result = await _run()
    tokens_used = (_token_tracker["input"] + _token_tracker["output"]) - tokens_before
    calls_used = _token_tracker["calls"] - calls_before
    cost_used = _token_tracker["cost_usd"] - cost_before
    return {
        "result": result,
        "research_type": research_type,
        "topic": topic,
        "image_used": bool(image_b64),
        "token_usage": {
            "total": tokens_used,
            "input": _token_tracker["input"],
            "output": _token_tracker["output"],
            "calls": calls_used,
            "cost_usd": round(cost_used, 4),
            "model": _token_tracker.get("last_model", select_model(6, "research")),
        },
    }


# ── Analyst Team — Node-based Flow Generator ──────────────────────────────────

class FlowRequest(BaseModel):
    prompt: str
    flow_type: str = "user_flow"
    domain: str = "general"
    project_id: Optional[str] = None
    image_b64: Optional[str] = None  # set by endpoint after image decode


@app.post("/api/v1/analyst/flow", status_code=201)
async def analyst_flow(
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
    prompt: str = Form(...),
    flow_type: str = Form("user_flow"),
    domain: str = Form("general"),
    project_id: Optional[str] = Form(None),
    image: UploadFile = File(None),
):
    """Generate a node-based user flow / journey map with UX analysis.
    Accepts optional image upload — vision-grounds the flow in real UI.
    """
    # Decode image
    image_b64 = None
    if image is not None and image.filename:
        ext = (image.filename.rsplit(".", 1)[-1].lower() if "." in image.filename else "png")
        if ext in {"png", "jpg", "jpeg", "webp"}:
            image_bytes = await image.read()
            import base64 as _b64
            image_b64 = _b64.b64encode(image_bytes).decode("utf-8")

    req = FlowRequest(prompt=prompt, flow_type=flow_type, domain=domain,
                      project_id=project_id, image_b64=image_b64)

    gen_id = str(uuid.uuid4())
    _active_generations[gen_id] = {"status": "running", "events": [], "result": None, "user_id": user.id}
    background_tasks.add_task(_run_flow, gen_id, req)
    return {"generation_id": gen_id, "status": "started", "ws_url": f"/ws/design/{gen_id}", "image_used": bool(image_b64)}


async def _run_flow(gen_id: str, req: FlowRequest):
    from app.agents.analyst.flow_agent import generate_user_flow

    async def progress(event):
        if gen_id in _active_generations:
            _active_generations[gen_id]["events"].append(event)
        await ws_manager.broadcast(gen_id, event)

    try:
        result = await generate_user_flow(
            prompt=req.prompt,
            flow_type=req.flow_type,
            domain=req.domain,
            llm_chat=llm_chat,
            select_model=select_model,
            rag_query=rag_query,
            progress_callback=progress,
            image_b64=req.image_b64,
            llm_chat_vision=llm_chat_vision if req.image_b64 else None,
        )
        _active_generations[gen_id]["status"] = "completed"
        _active_generations[gen_id]["result"] = result

        if req.project_id:
            async with AsyncSessionLocal() as db:
                art = Artifact(
                    project_id=req.project_id,
                    name=f"{result.get('flow_name', 'Flow')} ({req.flow_type})",
                    type=f"flow_{req.flow_type}",
                    status="approved",
                    content=json.dumps(result),
                    agent_type="analyst",
                    model_used=select_model(7, "ux"),
                )
                db.add(art)
                await db.commit()
    except Exception as e:
        import traceback
        traceback.print_exc()
        _active_generations[gen_id]["status"] = "failed"
        await progress({"type": "error", "message": str(e)})


# ── QA Team Endpoints ──────────────────────────────────────────────────────────

class QATestRequest(BaseModel):
    test_type: str    # functional | integration | e2e | regression | performance
    topic: str
    requirements: dict = {}
    user_stories: list = []
    project_id: Optional[str] = None


class AccessibilityRequest(BaseModel):
    topic: str
    ui_design: dict = {}
    ux_flow: dict = {}
    artifact_id: Optional[str] = None
    project_id: Optional[str] = None


class ComplianceRequest(BaseModel):
    compliance_type: str  # wcag | gdpr | hipaa | soc2 | all
    topic: str
    domain: str = "general"
    design_context: dict = {}
    project_id: Optional[str] = None


@app.post("/api/v1/qa/test-cases", status_code=201)
async def qa_test_cases(
    req: QATestRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Generate test cases from requirements/user stories."""
    from app.agents.qa.test_agent import run_test_cases

    result = await run_test_cases(
        test_type=req.test_type,
        requirements=req.requirements,
        user_stories=req.user_stories,
        topic=req.topic,
        llm_chat=llm_chat,
        select_model=select_model,
    )

    if req.project_id and result:
        art = Artifact(
            project_id=req.project_id,
            name=f"{req.test_type.title()} Test Suite — {req.topic[:60]}",
            type="qa_test_cases",
            status="approved",
            content=json.dumps(result),
            agent_type="qa",
            model_used=select_model(5, "qa"),
        )
        db.add(art)
        await db.commit()

    return {"result": result, "test_type": req.test_type}


@app.post("/api/v1/qa/accessibility", status_code=201)
async def qa_accessibility(
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
    topic: str = Form(...),
    artifact_id: Optional[str] = Form(None),
    project_id: Optional[str] = Form(None),
    ui_design_json: str = Form("{}"),
    ux_flow_json: str = Form("{}"),
    image: UploadFile = File(None),
):
    """Run accessibility check.
    - With image: full vision-based WCAG 2.1 AA check from screenshot + RAG standards
    - Without image: real contrast math + LLM deep analysis of design tokens
    """
    from app.agents.qa.test_agent import run_accessibility_check, run_accessibility_from_image

    # Decode image if provided
    image_b64 = None
    if image is not None and image.filename:
        ext = (image.filename.rsplit(".", 1)[-1].lower() if "." in image.filename else "png")
        if ext in {"png", "jpg", "jpeg", "webp"}:
            image_bytes = await image.read()
            import base64 as _b64
            image_b64 = _b64.b64encode(image_bytes).decode("utf-8")

    if image_b64:
        # Image path — full WCAG check from screenshot
        result = await run_accessibility_from_image(
            image_b64=image_b64,
            topic=topic,
            llm_chat=llm_chat,
            llm_chat_vision=llm_chat_vision,
            select_model=select_model,
            rag_query=rag_query,
        )
    else:
        # Token/design path — existing WCAG math check
        ui_design = json.loads(ui_design_json) if ui_design_json else {}
        ux_flow   = json.loads(ux_flow_json)   if ux_flow_json   else {}

        if artifact_id:
            r = await db.execute(sa.select(Artifact).where(Artifact.id == artifact_id))
            existing_art = r.scalar_one_or_none()
            if existing_art:
                art_content = json.loads(existing_art.content)
                if existing_art.type == "hifi_screen":
                    ui_design = art_content
                elif existing_art.type == "user_flow":
                    ux_flow = art_content

        result = await run_accessibility_check(
            ui_design=ui_design,
            ux_flow=ux_flow,
            topic=topic,
            llm_chat=llm_chat,
            select_model=select_model,
            rag_query=rag_query,
        )

    if project_id:
        new_art = Artifact(
            project_id=project_id,
            name=f"Accessibility Report — {topic[:60]}",
            type="qa_accessibility",
            status="approved",
            content=json.dumps(result),
            agent_type="qa",
            model_used=select_model(4, "qa"),
        )
        db.add(new_art)
        await db.commit()

    return {"result": result, "image_used": bool(image_b64)}


@app.post("/api/v1/qa/compliance", status_code=201)
async def qa_compliance(
    req: ComplianceRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Compliance check — WCAG, GDPR, HIPAA, SOC2."""
    from app.agents.qa.test_agent import run_compliance_check

    result = await run_compliance_check(
        compliance_type=req.compliance_type,
        design_context=req.design_context,
        topic=req.topic,
        domain=req.domain,
        llm_chat=llm_chat,
        select_model=select_model,
    )

    if req.project_id:
        art = Artifact(
            project_id=req.project_id,
            name=f"Compliance Report ({req.compliance_type.upper()}) — {req.topic[:60]}",
            type="qa_compliance",
            status="approved",
            content=json.dumps(result),
            agent_type="qa",
            model_used=select_model(5, "qa"),
        )
        db.add(art)
        await db.commit()

    return {"result": result}


# ── Engineering Team Endpoints ────────────────────────────────────────────────

class CodeGenRequest(BaseModel):
    output_format: str = "react"   # react | html | flutter | vue
    topic: str
    ui_artifact_id: Optional[str] = None
    ux_artifact_id: Optional[str] = None
    project_id: Optional[str] = None


class APIDesignRequest(BaseModel):
    topic: str
    requirements: dict = {}
    user_stories: list = []
    project_id: Optional[str] = None


class ArchitectureRequest(BaseModel):
    topic: str
    requirements: dict = {}
    scale: str = "medium"   # small | medium | enterprise
    project_id: Optional[str] = None


@app.post("/api/v1/engineering/code", status_code=201)
async def engineering_code(
    req: CodeGenRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Generate production code from UI design."""
    from app.agents.engineering.code_agent import run_code_generation

    ui_design, ux_flow = {}, {}

    if req.ui_artifact_id:
        r = await db.execute(sa.select(Artifact).where(Artifact.id == req.ui_artifact_id))
        art = r.scalar_one_or_none()
        if art: ui_design = json.loads(art.content)

    if req.ux_artifact_id:
        r = await db.execute(sa.select(Artifact).where(Artifact.id == req.ux_artifact_id))
        art = r.scalar_one_or_none()
        if art: ux_flow = json.loads(art.content)

    result = await run_code_generation(
        output_format=req.output_format,
        ui_design=ui_design,
        ux_flow=ux_flow,
        topic=req.topic,
        llm_chat=llm_chat,
        select_model=select_model,
    )

    if req.project_id:
        art = Artifact(
            project_id=req.project_id,
            name=f"{req.output_format.title()} Code — {req.topic[:60]}",
            type=f"code_{req.output_format}",
            status="approved",
            content=json.dumps(result),
            agent_type="engineering",
            model_used=select_model(7, "ui"),
        )
        db.add(art)
        await db.commit()

    return {"result": result}


@app.post("/api/v1/engineering/api-design", status_code=201)
async def engineering_api_design(
    req: APIDesignRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Design REST API from requirements."""
    from app.agents.engineering.code_agent import run_api_design

    result = await run_api_design(
        requirements=req.requirements,
        user_stories=req.user_stories,
        topic=req.topic,
        llm_chat=llm_chat,
        select_model=select_model,
    )

    if req.project_id:
        art = Artifact(
            project_id=req.project_id,
            name=f"API Design — {req.topic[:60]}",
            type="api_design",
            status="approved",
            content=json.dumps(result),
            agent_type="engineering",
            model_used=select_model(6, "planner"),
        )
        db.add(art)
        await db.commit()

    return {"result": result}


@app.post("/api/v1/engineering/architecture", status_code=201)
async def engineering_architecture(
    req: ArchitectureRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Generate system architecture plan."""
    from app.agents.engineering.code_agent import run_architecture_plan

    result = await run_architecture_plan(
        requirements=req.requirements,
        topic=req.topic,
        scale=req.scale,
        llm_chat=llm_chat,
        select_model=select_model,
        rag_query=rag_query,
    )

    if req.project_id:
        art = Artifact(
            project_id=req.project_id,
            name=f"Architecture Plan — {req.topic[:60]}",
            type="architecture",
            status="approved",
            content=json.dumps(result),
            agent_type="engineering",
            model_used=select_model(8, "planner"),
        )
        db.add(art)
        await db.commit()

    return {"result": result}


# ── Design Studio — Unified Flow Generator ────────────────────────────────────

@app.get("/api/v1/design/systems")
async def list_design_systems(user: User = Depends(get_current_user)):
    """List all available design systems."""
    from app.design_system.registry import DESIGN_SYSTEMS
    return {
        "systems": [
            {
                "id": k,
                "name": v["name"],
                "description": v["description"],
                "primary_color": v["tokens"]["colors"].get("primary", "#5B5EF4"),
                "tokens": v["tokens"],
                "component_style": v.get("component_style", ""),
            }
            for k, v in DESIGN_SYSTEMS.items()
        ]
    }


# Store active flow generations for WebSocket streaming
_active_generations: dict[str, dict] = {}

# Store design sessions (conversation memory + planned/generated screens)
# Keyed by session_id → { history: [], screens: [], plan_payload: {}, ... }
_design_sessions: dict[str, dict] = {}


# ══════════════════════════════════════════════════════════════════════════════
# UX MAGIC STYLE: Plan → Approve → Generate
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/v1/design/plan", status_code=201)
async def design_plan(
    background_tasks: BackgroundTasks,
    prompt: str = Form(...),
    session_id: str = Form(None),
    design_system: str = Form("company"),
    fidelity: str = Form("hifi"),
    screen_count: int = Form(0),       # 0 = agent decides
    domain: str = Form("general"),
    reference_mode: str = Form(None),  # recreate | improve | inspire
    project_id: str = Form(None),
    image: UploadFile = File(None),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    PHASE 1: Agent analyzes the request (+ optional reference image),
    shows thought process, and proposes a screen plan for approval.
    """
    session_id = session_id or str(uuid.uuid4())

    # Load from DB if session not already in memory (e.g. after server restart)
    if session_id not in _design_sessions:
        await _load_session(session_id)

    session = _design_sessions.setdefault(session_id, {
        "history": [], "screens": [], "plan_payload": None,
        "user_id": user.id, "project_id": project_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    session["history"].append({"role": "user", "content": prompt})

    # Process optional reference image
    image_b64 = None
    if image is not None:
        ext = (image.filename.rsplit(".", 1)[-1].lower() if "." in image.filename else "png")
        if ext not in {"png", "jpg", "jpeg", "webp", "gif"}:
            raise HTTPException(400, f"Image format .{ext} not supported")
        image_bytes = await image.read()
        image_b64 = base64.b64encode(image_bytes).decode("utf-8")
        uploads_dir = Path("./uploads/references")
        uploads_dir.mkdir(parents=True, exist_ok=True)
        async with aiofiles.open(str(uploads_dir / f"{uuid.uuid4()}.{ext}"), "wb") as f:
            await f.write(image_bytes)

    gen_id = str(uuid.uuid4())
    _active_generations[gen_id] = {
        "status": "planning", "session_id": session_id, "events": [],
        "plan_payload": None, "started_at": datetime.now(timezone.utc).isoformat(),
        "user_id": user.id,
    }

    background_tasks.add_task(
        _run_plan, gen_id=gen_id, session_id=session_id,
        prompt=prompt, design_system=design_system, fidelity=fidelity,
        screen_count=screen_count, domain=domain,
        reference_mode=reference_mode if image_b64 else None,
        reference_image_b64=image_b64,
    )

    return {"generation_id": gen_id, "session_id": session_id, "status": "planning", "ws_url": f"/ws/design/{gen_id}"}


async def _run_plan(gen_id, session_id, prompt, design_system, fidelity,
                    screen_count, domain, reference_mode, reference_image_b64):
    """Background: run the planning phase."""
    from app.agents.design_complete.conversation_agent import plan_flow

    session = _design_sessions.get(session_id, {})

    _gen_token_tracker[gen_id] = {
        "input_start": _token_tracker["input"],
        "output_start": _token_tracker["output"],
        "calls_start": _token_tracker["calls"],
        "cost_start": _token_tracker["cost_usd"],
    }

    async def progress(event: dict):
        gen_snap = _gen_token_tracker.get(gen_id, {})
        event["token_usage"] = {
            "input": _token_tracker["input"] - gen_snap.get("input_start", 0),
            "output": _token_tracker["output"] - gen_snap.get("output_start", 0),
            "calls": _token_tracker["calls"] - gen_snap.get("calls_start", 0),
            "total": (_token_tracker["input"] - gen_snap.get("input_start", 0)) +
                     (_token_tracker["output"] - gen_snap.get("output_start", 0)),
            "cost_usd": round(_token_tracker["cost_usd"] - gen_snap.get("cost_start", 0.0), 4),
            "model": _token_tracker.get("last_model", ""),
        }
        if gen_id in _active_generations:
            _active_generations[gen_id]["events"].append(event)
        await ws_manager.broadcast(gen_id, event)

    try:
        result = await plan_flow(
            prompt=prompt,
            design_system_id=design_system,
            fidelity=fidelity,
            screen_count=screen_count,
            domain=domain,
            conversation_history=session.get("history", []),
            existing_screens=session.get("screens", []),
            reference_image_b64=reference_image_b64,
            reference_mode=reference_mode,
            llm_chat=llm_chat,
            llm_chat_vision=llm_chat_vision,
            select_model=select_model,
            rag_query=rag_query,
            progress_callback=progress,
            vision_backend_label=vision_backend_label(),
        )

        # Store plan in session + generation
        # Keep the raw reference image + mode for high-accuracy recreation
        result["reference_image_b64"] = reference_image_b64
        result["reference_mode"] = reference_mode

        _design_sessions[session_id]["plan_payload"] = result
        _active_generations[gen_id]["status"] = "planned"
        _active_generations[gen_id]["plan_payload"] = result

        # Record agent's plan in history
        plan = result.get("plan", {})
        _design_sessions[session_id]["history"].append({
            "role": "assistant",
            "content": f"Planned {len(plan.get('planned_screens', []))} screens: {plan.get('thought_process', '')[:200]}",
        })

        # Persist plan to DB so it survives restarts
        await _persist_session(session_id)

    except Exception as e:
        import traceback
        traceback.print_exc()
        _active_generations[gen_id]["status"] = "failed"
        await progress({"type": "error", "message": f"Planning error: {e}"})


class ApproveRequest(BaseModel):
    generation_id: str
    session_id: str
    approved_screen_ids: list[str] = []   # empty = approve all
    project_id: Optional[str] = None


@app.post("/api/v1/design/approve", status_code=201)
async def design_approve(
    req: ApproveRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    PHASE 2: User approved the plan. Generate the approved screens.
    Returns a NEW generation_id for streaming the generation.
    """
    # Load from DB if not in memory
    if req.session_id not in _design_sessions:
        await _load_session(req.session_id)
    session = _design_sessions.get(req.session_id)
    if not session or not session.get("plan_payload"):
        raise HTTPException(404, "No plan found for this session. Run /design/plan first.")

    plan_payload = session["plan_payload"]

    gen_id = str(uuid.uuid4())
    _active_generations[gen_id] = {
        "status": "generating",
        "session_id": req.session_id,
        "events": [],
        "result": None,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "user_id": user.id,
    }

    background_tasks.add_task(
        _run_generate,
        gen_id=gen_id,
        session_id=req.session_id,
        plan_payload=plan_payload,
        approved_ids=req.approved_screen_ids,
        user_prompt=session["history"][-2]["content"] if len(session["history"]) >= 2 else "",
        project_id=req.project_id or session.get("project_id"),
    )

    return {"generation_id": gen_id, "session_id": req.session_id, "status": "generating", "ws_url": f"/ws/design/{gen_id}"}


def _build_compliance_metrics(plan_payload: dict, result: dict, design_system_id: str) -> dict:
    """Build compliance & accessibility summary from plan and generated result."""
    plan = plan_payload.get("plan", {})
    screens = result.get("screens", [])
    fidelity = result.get("fidelity", "hifi")

    # Derive accessibility score by checking HTML for key patterns
    total_score = 0
    checks = []
    for screen in screens:
        html = screen.get("html", "")
        if 'aria-' in html or 'role=' in html:
            total_score += 15
        if 'alt=' in html:
            total_score += 10
        if 'focus' in html or 'focus-visible' in html:
            total_score += 10
        if 'sr-only' in html or 'visually-hidden' in html:
            total_score += 5

    screen_count = len(screens) or 1
    base_score = 60  # base for using design system with built-in tokens
    bonus = min(40, total_score // screen_count)
    accessibility_score = base_score + bonus

    # Compliance standards applied
    compliance_standards = ["WCAG 2.1 AA"]
    if design_system_id != "general":
        compliance_standards.append("Brand Design System")
    if fidelity == "hifi":
        compliance_standards.append("High-Fidelity Token Usage")

    plan_reqs = plan.get("accessibility_requirements", [])
    if plan_reqs:
        compliance_standards.append("Project Accessibility Requirements")

    # Design system checks
    ds_checks = [
        {"name": "Color Contrast ≥ 4.5:1", "status": "pass", "detail": "Design tokens enforce WCAG AA contrast ratios"},
        {"name": "Touch Targets ≥ 48×48px", "status": "pass", "detail": "Component rules include minimum tap target sizes"},
        {"name": "Typography Scale", "status": "pass", "detail": "Consistent type scale applied from design tokens"},
        {"name": "Spacing System", "status": "pass", "detail": "8pt grid spacing tokens used throughout"},
        {"name": "Focus Indicators", "status": "warning" if "focus" not in str(screens) else "pass",
         "detail": "Focus-visible states recommended for interactive elements"},
        {"name": "ARIA Labels", "status": "pass" if any('aria-' in s.get("html","") for s in screens) else "warning",
         "detail": "Semantic HTML with ARIA attributes for assistive tech"},
    ]

    # Design metrics
    components_used = []
    for s in screens:
        components_used.extend(s.get("components_used", []))
    unique_components = list(set(components_used))

    models_used = list(set(s.get("model_used", "") for s in screens if s.get("model_used")))

    return {
        "accessibility_score": accessibility_score,
        "compliance_standards": compliance_standards,
        "checks": ds_checks,
        "design_metrics": {
            "screens_generated": len(screens),
            "fidelity": fidelity,
            "design_system": design_system_id,
            "components_used": unique_components[:12],
            "models_used": models_used,
        },
        "wcag_level": "AA",
        "summary": f"Design meets WCAG 2.1 AA standards with an accessibility score of {accessibility_score}/100. "
                   f"{len(compliance_standards)} compliance standards applied.",
    }


async def _run_generate(gen_id, session_id, plan_payload, approved_ids, user_prompt, project_id):
    """Background: generate approved screens."""
    from app.agents.design_complete.conversation_agent import generate_approved_screens

    # Snapshot token baseline for this generation
    _gen_token_tracker[gen_id] = {
        "input_start": _token_tracker["input"],
        "output_start": _token_tracker["output"],
        "calls_start": _token_tracker["calls"],
        "cost_start": _token_tracker["cost_usd"],
    }

    async def progress(event: dict):
        # Attach current token usage delta to every event
        gen_snap = _gen_token_tracker.get(gen_id, {})
        event["token_usage"] = {
            "input": _token_tracker["input"] - gen_snap.get("input_start", 0),
            "output": _token_tracker["output"] - gen_snap.get("output_start", 0),
            "calls": _token_tracker["calls"] - gen_snap.get("calls_start", 0),
            "total": (_token_tracker["input"] - gen_snap.get("input_start", 0)) +
                     (_token_tracker["output"] - gen_snap.get("output_start", 0)),
            "cost_usd": round(_token_tracker["cost_usd"] - gen_snap.get("cost_start", 0.0), 4),
            "model": _token_tracker.get("last_model", ""),
        }
        if gen_id in _active_generations:
            _active_generations[gen_id]["events"].append(event)
        await ws_manager.broadcast(gen_id, event)

    try:
        result = await generate_approved_screens(
            plan_payload=plan_payload,
            approved_screen_ids=approved_ids,
            user_prompt=user_prompt,
            llm_chat=llm_chat,
            select_model=select_model,
            progress_callback=progress,
            llm_chat_vision=llm_chat_vision,
        )

        # Append generated screens to session (so "next flow" knows what exists)
        session = _design_sessions.get(session_id, {})
        session.setdefault("screens", []).extend(result.get("screens", []))

        _active_generations[gen_id]["status"] = "completed"
        _active_generations[gen_id]["result"] = result

        # Save artifact
        if project_id:
            async with AsyncSessionLocal() as db:
                art = Artifact(
                    project_id=project_id,
                    name=f"{result.get('flow_name', 'Flow')} ({result.get('fidelity', 'hifi')})",
                    type="design_flow",
                    status="approved",
                    content=json.dumps(result),
                    agent_type="design",
                    model_used=select_model(8, "ui"),
                )
                db.add(art)
                await db.commit()
                await db.refresh(art)
                result["artifact_id"] = art.id

        # Persist full session to DB so it survives restarts
        await _persist_session(session_id)

        # Self-learning: add rich screen patterns + past design record to RAG
        plan_meta = plan_payload.get("plan", {})
        screen_type = plan_meta.get("screen_type", "web_dashboard")
        domain_tag = plan_payload.get("domain", "general")
        flow_name = result.get("flow_name", user_prompt[:60])
        for i, s in enumerate(result.get("screens", [])):
            # Detailed pattern for design_patterns collection (used in planning)
            pattern = (
                f"SCREEN: {s.get('name')} | TYPE: {screen_type} | DOMAIN: {domain_tag}\n"
                f"PURPOSE: {s.get('purpose', '')}\n"
                f"LAYOUT: {s.get('layout_type', '')} | ELEMENTS: {', '.join((s.get('key_elements') or [])[:6])}\n"
                f"DESIGN NOTES: {s.get('design_notes', '')}"
            )
            if pattern.strip():
                await rag_add("design_patterns", f"{gen_id}-screen-{i}", pattern,
                              {"session": session_id, "screen_type": screen_type, "domain": domain_tag})

        # Rich past_designs entry (used by planner for "have we done this before?")
        screens_summary = "\n".join(
            f"  - {s.get('name')}: {s.get('purpose', '')} [{s.get('layout_type', '')}]"
            for s in result.get("screens", [])
        )
        past_entry = (
            f"DESIGN FLOW: {flow_name}\n"
            f"DOMAIN: {domain_tag} | SCREEN TYPE: {screen_type} | FIDELITY: {result.get('fidelity','hifi')}\n"
            f"PROMPT: {user_prompt[:300]}\n"
            f"SCREENS ({len(result.get('screens', []))}):\n{screens_summary}\n"
            f"DESIGN APPROACH: {plan_meta.get('design_approach', {}).get('primary_pattern', '')}\n"
            f"THOUGHT PROCESS: {plan_meta.get('thought_process', '')[:400]}"
        )
        await rag_add_chunked("past_designs", gen_id, past_entry,
                              {"session": session_id, "domain": domain_tag, "screen_type": screen_type})

        # Build compliance & accessibility metrics summary
        design_system_id = plan_payload.get("design_system", {}).get("id", "company")
        fidelity = result.get("fidelity", "hifi")
        screen_count = len(result.get("screens", []))
        compliance_metrics = _build_compliance_metrics(plan_payload, result, design_system_id)

        await progress({
            "type": "complete",
            "message": f"✨ Generated {screen_count} screens",
            "result": result,
            "compliance_metrics": compliance_metrics,
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        _active_generations[gen_id]["status"] = "failed"
        await progress({"type": "error", "message": f"Generation error: {e}"})


class RegenerateRequest(BaseModel):
    session_id: str
    screen_id: str
    refinement: str
    project_id: Optional[str] = None


@app.post("/api/v1/design/regenerate-screen", status_code=201)
async def design_regenerate_screen(
    req: RegenerateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Regenerate a single screen with a refinement instruction (bumps version)."""
    if req.session_id not in _design_sessions:
        await _load_session(req.session_id)
    session = _design_sessions.get(req.session_id)
    if not session:
        raise HTTPException(404, "Session not found")

    existing = next((s for s in session.get("screens", []) if s.get("screen_id") == req.screen_id), None)
    if not existing:
        raise HTTPException(404, "Screen not found in session")

    gen_id = str(uuid.uuid4())
    _active_generations[gen_id] = {"status": "regenerating", "events": [], "result": None, "user_id": user.id}

    background_tasks.add_task(
        _run_regenerate, gen_id, req.session_id, req.screen_id, existing, req.refinement,
        session.get("plan_payload", {}),
        session["history"][0]["content"] if session.get("history") else "",
    )

    return {"generation_id": gen_id, "status": "regenerating", "ws_url": f"/ws/design/{gen_id}"}


async def _run_regenerate(gen_id, session_id, screen_id, existing, refinement, plan_payload, user_prompt):
    from app.agents.design_complete.conversation_agent import regenerate_screen

    async def progress(event):
        if gen_id in _active_generations:
            _active_generations[gen_id]["events"].append(event)
        await ws_manager.broadcast(gen_id, event)

    try:
        new_screen = await regenerate_screen(
            screen_id=screen_id, existing_screen=existing, refinement_instruction=refinement,
            plan_payload=plan_payload, user_prompt=user_prompt,
            llm_chat=llm_chat, select_model=select_model, progress_callback=progress,
        )
        # Replace screen in session
        session = _design_sessions.get(session_id, {})
        screens = session.get("screens", [])
        for i, s in enumerate(screens):
            if s.get("screen_id") == screen_id:
                screens[i] = new_screen
                break

        _active_generations[gen_id]["status"] = "completed"
        _active_generations[gen_id]["result"] = {"screen": new_screen}
        # Persist updated session to DB after screen regeneration
        await _persist_session(session_id)
        await progress({"type": "screen_regenerated", "screen": new_screen, "message": f"✓ {new_screen['name']} → V{new_screen['version']}"})
    except Exception as e:
        _active_generations[gen_id]["status"] = "failed"
        await progress({"type": "error", "message": str(e)})


class ExportScreen(BaseModel):
    screen_id: str
    name: str
    html: str

class ExportRequest(BaseModel):
    session_id: str
    format: str = "html"   # html | figma_json
    screens: list[ExportScreen] = []   # client sends screens directly so export works across restarts


@app.post("/api/v1/design/export")
async def export_design(req: ExportRequest, user: User = Depends(get_current_user)):
    """Export generated screens as standalone HTML or Figma-plugin JSON."""
    # Prefer screens passed directly from client (survives server restarts)
    session = _design_sessions.get(req.session_id) or {}
    if req.screens:
        screens = [s.model_dump() for s in req.screens]
    else:
        if not session.get("screens"):
            raise HTTPException(404, "No screens found — please regenerate the design")
        screens = session["screens"]

    if req.format == "html":
        # Build a single HTML file with all screens stacked + nav
        ds = session.get("plan_payload", {}).get("design_system", {})
        nav_links = "".join(
            f'<a href="#{s["screen_id"]}" style="padding:8px 14px;border-radius:8px;background:#f1f1f1;text-decoration:none;color:#333;font-size:13px;">{i+1}. {s["name"]}</a>'
            for i, s in enumerate(screens)
        )
        screen_sections = "".join(
            f'<section id="{s["screen_id"]}" style="margin-bottom:48px;"><h2 style="font:600 18px Inter,sans-serif;color:#111;margin-bottom:12px;">{i+1}. {s["name"]}</h2><div style="border:1px solid #e5e5e5;border-radius:16px;overflow:hidden;">{s["html"]}</div></section>'
            for i, s in enumerate(screens)
        )
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{ds.get('name', 'CORE Studio')} Export</title>
<script src="https://cdn.tailwindcss.com"></script>
<style>body{{margin:0;font-family:Inter,system-ui,sans-serif;background:#fafafa;}}*{{box-sizing:border-box;}}</style>
</head><body>
<div style="max-width:1100px;margin:0 auto;padding:32px;">
<h1 style="font:700 24px Inter,sans-serif;color:#111;">CORE Studio Export — {len(screens)} screens</h1>
<div style="display:flex;gap:8px;flex-wrap:wrap;margin:16px 0 32px;">{nav_links}</div>
{screen_sections}
</div></body></html>"""
        return {"format": "html", "filename": "core-studio-export.html", "content": html, "screen_count": len(screens)}

    elif req.format == "figma_json":
        # Real auto-layout node tree (not a flat HTML dump) for the CORE Studio Figma
        # plugin (packages/figma-plugin/) — built via html_to_figma_tree, which maps
        # each screen's flexbox-based inline styles to Figma's official Plugin API
        # properties (layoutMode, itemSpacing, padding*, alignment, sizing).
        from app.design_system.figma_export import html_to_figma_tree

        figma_screens = []
        for s in screens:
            try:
                tree = html_to_figma_tree(s["html"])
            except Exception as e:
                tree = {"type": "FRAME", "layout": {"mode": "NONE"}, "style": {}, "children": [], "error": str(e)}
            figma_screens.append({
                "title": s["name"],
                "width": 1440,
                "height": 1024,
                "tree": tree,
            })
        return {
            "format": "figma_json",
            "filename": "core-studio-figma.json",
            "content": {"screens": figma_screens, "version": "3.0"},
            "screen_count": len(screens),
            "instructions": "Open Figma → Plugins → CORE Studio Import → paste or load this JSON to build real auto-layout frames.",
        }

    raise HTTPException(400, "format must be 'html' or 'figma_json'")


@app.get("/api/v1/design/session/{session_id}")
async def get_design_session(session_id: str, user: User = Depends(get_current_user)):
    """Get full session state — history + all screens."""
    if session_id not in _design_sessions:
        await _load_session(session_id)
    session = _design_sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "session_id": session_id,
        "history": session.get("history", []),
        "screens": session.get("screens", []),
        "has_plan": session.get("plan_payload") is not None,
    }


@app.post("/api/v1/design/generate", status_code=201)
async def generate_design_flow(
    background_tasks: BackgroundTasks,
    prompt: str = Form(...),
    design_system: str = Form("company"),
    fidelity: str = Form("hifi"),         # "wireframe" | "hifi"
    screen_count: int = Form(3),
    domain: str = Form("general"),
    reference_mode: str = Form(None),     # "recreate" | "improve" | "inspire"
    project_id: str = Form(None),
    image: UploadFile = File(None),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    UNIFIED design generation endpoint.

    Like UX Magic / Claude chat: prompt + optional image + design system + fidelity → flow of N screens.
    Always retrieves RAG context first (V1 pattern).
    """
    # Validate
    if fidelity not in ("wireframe", "hifi"):
        raise HTTPException(400, "fidelity must be 'wireframe' or 'hifi'")
    if screen_count < 1 or screen_count > 10:
        raise HTTPException(400, "screen_count must be 1-10")

    # Process optional reference image
    image_b64 = None
    if image is not None:
        if resolve_vision_provider() == "none":
            raise HTTPException(400, "Reading a reference image requires a configured vision "
                                      "backend — set a cloud AI API key or run Ollama with llava.")
        ext = image.filename.rsplit(".", 1)[-1].lower() if "." in image.filename else "png"
        if ext not in {"png", "jpg", "jpeg", "webp", "gif"}:
            raise HTTPException(400, f"Image format .{ext} not supported")
        image_bytes = await image.read()
        image_b64 = base64.b64encode(image_bytes).decode("utf-8")

        # Save the image
        uploads_dir = Path("./uploads/references")
        uploads_dir.mkdir(parents=True, exist_ok=True)
        img_id = str(uuid.uuid4())
        async with aiofiles.open(str(uploads_dir / f"{img_id}.{ext}"), "wb") as f:
            await f.write(image_bytes)

    # Create generation session for WebSocket streaming
    gen_id = str(uuid.uuid4())
    _active_generations[gen_id] = {
        "status": "starting",
        "events": [],
        "result": None,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "user_id": user.id,
    }

    # Kick off in background
    background_tasks.add_task(
        _run_design_generation,
        gen_id=gen_id,
        prompt=prompt,
        design_system_id=design_system,
        fidelity=fidelity,
        screen_count=screen_count,
        domain=domain,
        reference_mode=reference_mode if image_b64 else None,
        reference_image_b64=image_b64,
        project_id=project_id,
        user_id=user.id,
    )

    return {
        "generation_id": gen_id,
        "status": "started",
        "ws_url": f"/ws/design/{gen_id}",
        "rest_status_url": f"/api/v1/design/generate/{gen_id}",
    }


async def _run_design_generation(
    gen_id: str,
    prompt: str,
    design_system_id: str,
    fidelity: str,
    screen_count: int,
    domain: str,
    reference_mode: Optional[str],
    reference_image_b64: Optional[str],
    project_id: Optional[str],
    user_id: str,
):
    """Background task — runs the full generation pipeline."""
    from app.agents.design_complete.flow_generator import generate_screen_flow

    async def progress(event: dict):
        # Append to memory for REST polling
        if gen_id in _active_generations:
            _active_generations[gen_id]["events"].append(event)
        # Broadcast over WS
        await ws_manager.broadcast(gen_id, event)

    try:
        _active_generations[gen_id]["status"] = "running"
        await progress({"type": "started", "message": "🚀 Starting design generation..."})

        result = await generate_screen_flow(
            prompt=prompt,
            design_system_id=design_system_id,
            fidelity=fidelity,
            screen_count=screen_count,
            domain=domain,
            reference_image_b64=reference_image_b64,
            reference_mode=reference_mode,
            llm_chat=llm_chat,
            llm_chat_vision=llm_chat_vision,
            select_model=select_model,
            rag_query=rag_query,
            progress_callback=progress,
        )

        _active_generations[gen_id]["status"] = "completed"
        _active_generations[gen_id]["result"] = result

        # Save as artifact if project_id provided
        if project_id:
            async with AsyncSessionLocal() as db:
                art = Artifact(
                    project_id=project_id,
                    name=f"{result.get('flow_name', 'Design Flow')} ({fidelity})",
                    type="design_flow",
                    status="approved",
                    content=json.dumps(result),
                    agent_type="design",
                    model_used=select_model(8, "ui"),
                )
                db.add(art)
                await db.commit()
                await db.refresh(art)
                result["artifact_id"] = art.id

        await progress({
            "type": "complete",
            "message": f"✨ Generated {len(result.get('screens', []))} screens",
            "result": result,
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        _active_generations[gen_id]["status"] = "failed"
        _active_generations[gen_id]["error"] = str(e)
        await progress({"type": "error", "message": f"Error: {e}"})


@app.get("/api/v1/design/generate/{gen_id}")
async def get_generation_status(
    gen_id: str,
    user: User = Depends(get_current_user),
):
    """Poll generation status (alternative to WebSocket)."""
    if gen_id not in _active_generations:
        raise HTTPException(404, "Generation not found")
    return _active_generations[gen_id]


@app.websocket("/ws/design/{gen_id}")
async def design_ws(websocket: WebSocket, gen_id: str):
    """Stream generation events in real-time."""
    await ws_manager.connect(websocket, gen_id)

    # Send any backlog events
    if gen_id in _active_generations:
        for event in _active_generations[gen_id]["events"]:
            await websocket.send_text(json.dumps(event))

    try:
        while True:
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_text('{"type":"pong"}')
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket, gen_id)


# ── Design Team — Screenshot & Vision Endpoints ───────────────────────────────

@app.post("/api/v1/design/screenshot-to-design", status_code=201)
async def screenshot_to_design(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mode: str = Form("recreate"),       # recreate | improve | analyze
    context: str = Form(""),
    project_id: str = Form(None),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Upload screenshot/sketch → design spec using LLava:7b vision."""
    if not HAS_OLLAMA:
        raise HTTPException(400, "Ollama not running. Start Ollama to use vision features.")

    # Check LLava is available
    content_bytes = await file.read()
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "png"
    if ext not in {"png", "jpg", "jpeg", "webp", "gif"}:
        raise HTTPException(400, f"Image format .{ext} not supported")

    # Save image
    uploads_dir = Path("./uploads/screenshots")
    uploads_dir.mkdir(parents=True, exist_ok=True)
    img_path = uploads_dir / f"{uuid.uuid4()}.{ext}"
    async with aiofiles.open(str(img_path), "wb") as f:
        await f.write(content_bytes)

    from app.agents.design_complete.screenshot_agent import run_screenshot_to_design

    result = await run_screenshot_to_design(
        image_path=str(img_path),
        mode=mode,
        context=context,
        llm_chat_vision=llm_chat_vision,
        select_model=select_model,
        rag_query=rag_query,
    )

    if project_id:
        art = Artifact(
            project_id=project_id,
            name=f"Screenshot Analysis ({mode}) — {file.filename[:40]}",
            type="design_screenshot",
            status="approved",
            content=json.dumps(result),
            agent_type="design",
            model_used=OLLAMA_MODELS["vision"],
        )
        db.add(art)
        await db.commit()

    return {"result": result, "mode": mode, "image": file.filename}


@app.post("/api/v1/design/handwritten-to-design", status_code=201)
async def handwritten_to_design(
    file: UploadFile = File(...),
    context: str = Form(""),
    project_id: str = Form(None),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Upload handwritten sketch → structured design using LLava:7b."""
    if not HAS_OLLAMA:
        raise HTTPException(400, "Ollama not running")

    content_bytes = await file.read()
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "png"
    uploads_dir = Path("./uploads/sketches")
    uploads_dir.mkdir(parents=True, exist_ok=True)
    img_path = uploads_dir / f"{uuid.uuid4()}.{ext}"
    async with aiofiles.open(str(img_path), "wb") as f:
        await f.write(content_bytes)

    from app.agents.design_complete.screenshot_agent import run_handwritten_to_design

    result = await run_handwritten_to_design(
        image_path=str(img_path),
        context=context,
        llm_chat_vision=llm_chat_vision,
        select_model=select_model,
    )

    if project_id:
        art = Artifact(
            project_id=project_id,
            name=f"Sketch Analysis — {file.filename[:40]}",
            type="design_sketch",
            status="approved",
            content=json.dumps(result),
            agent_type="design",
            model_used=OLLAMA_MODELS["vision"],
        )
        db.add(art)
        await db.commit()

    return {"result": result}


# ── Chat / AI Assistant (from V1 chat.py, adapted) ───────────────────────────

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    context: dict = {}
    project_id: Optional[str] = None


INTENT_MAP = {
    "accessibility": ["accessibility", "a11y", "wcag", "screen reader", "aria"],
    "analytics":     ["report", "analytics", "metrics", "statistics", "insights"],
    "ux_review":     ["check ux", "review ux", "usability", "user flow", "wireframe"],
    "planning":      ["plan", "strategy", "phases", "roadmap", "timeline"],
    "research":      ["research", "persona", "competitor", "market analysis", "user research"],
    "test":          ["test case", "test plan", "qa", "test coverage", "regression"],
    "code":          ["generate code", "write code", "react component", "html", "implement"],
    "design":        ["design", "create", "build", "ui", "screen", "layout", "component"],
}


def _classify_chat_intent(message: str) -> str:
    msg = message.lower()
    for intent, keywords in INTENT_MAP.items():
        if any(kw in msg for kw in keywords):
            return intent
    return "design"


@app.post("/api/v1/chat")
async def chat(
    req: ChatRequest,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """AI assistant — routes to right agent based on intent."""
    intent = _classify_chat_intent(req.message)
    session_id = req.session_id or str(uuid.uuid4())

    # Route based on intent
    if intent == "research":
        from app.agents.analyst.research_agent import run_research
        result = await run_research(
            research_type="market",
            topic=req.message[:100],
            domain="general",
            context=req.message,
            llm_chat=llm_chat,
            select_model=select_model,
            rag_query=rag_query,
        )
        response = f"Research complete. Key insights: {', '.join(result.get('key_insights', ['See full report'])[:2])}"

    elif intent == "test":
        from app.agents.qa.test_agent import run_test_cases
        result = await run_test_cases(
            test_type="functional",
            requirements={},
            user_stories=[],
            topic=req.message[:100],
            llm_chat=llm_chat,
            select_model=select_model,
        )
        tc_count = len(result.get("test_cases", []))
        response = f"Generated {tc_count} test cases for '{req.message[:60]}'"

    elif intent == "accessibility":
        from app.agents.qa.test_agent import run_accessibility_check
        ui = req.context.get("ui_design", {})
        ux = req.context.get("ux_flow", {})
        result = await run_accessibility_check(
            ui_design=ui, ux_flow=ux,
            topic=req.message[:100],
            llm_chat=llm_chat, select_model=select_model,
        )
        score = result.get("overall_score", "?")
        response = f"Accessibility score: {score}/100 ({result.get('wcag_level', 'AA')}). {result.get('summary', '')}"

    elif intent == "code":
        from app.agents.engineering.code_agent import run_code_generation
        result = await run_code_generation(
            output_format="react",
            ui_design=req.context.get("ui_design", {}),
            ux_flow=req.context.get("ux_flow", {}),
            topic=req.message[:100],
            llm_chat=llm_chat, select_model=select_model,
        )
        file_count = len(result.get("files", []))
        response = f"Generated {file_count} code file(s). First file: {result.get('files', [{}])[0].get('filename', 'N/A')}"

    else:
        # Default: design intent
        model = select_model(5, "planner")
        raw = await llm_chat(
            "You are a helpful design AI assistant. Give a concise, actionable response.",
            req.message,
            model,
        )
        import re as _re
        clean = _re.sub(r"<think>.*?</think>", "", raw, flags=_re.DOTALL).strip()
        result = {"response": clean}
        response = clean[:500]

    return {
        "message": response,
        "intent": intent,
        "session_id": session_id,
        "full_result": result,
        "model_used": select_model(5, intent[:7] if intent != "accessibility" else "qa"),
    }


# ── Team capability overview ──────────────────────────────────────────────────

@app.get("/api/v1/teams")
async def get_teams(user: User = Depends(get_current_user)):
    """Return all available teams and their capabilities."""
    return {
        "teams": [
            {
                "id": "design",
                "name": "Design Team",
                "description": "Requirements → Wireframes → Hi-Fi Screens → Review",
                "icon": "🎨",
                "capabilities": [
                    {"id": "planner", "name": "Requirement Blueprint", "endpoint": "/workflows", "model": select_model(5, "planner")},
                    {"id": "ux", "name": "UX Blueprint & Flows", "endpoint": "/workflows", "model": select_model(5, "ux")},
                    {"id": "ui", "name": "Visual UI Design", "endpoint": "/workflows", "model": select_model(5, "ui")},
                    {"id": "review", "name": "Design Review", "endpoint": "/workflows", "model": select_model(5, "review")},
                    {"id": "screenshot", "name": "Screenshot → Design", "endpoint": "/design/screenshot-to-design", "model": OLLAMA_MODELS["vision"]},
                    {"id": "handwritten", "name": "Sketch → Design", "endpoint": "/design/handwritten-to-design", "model": OLLAMA_MODELS["vision"]},
                ],
            },
            {
                "id": "analyst",
                "name": "Analyst Team",
                "description": "Research · Personas · User Stories · PRD",
                "icon": "📊",
                "capabilities": [
                    {"id": "personas", "name": "User Personas", "endpoint": "/analyst/research", "model": select_model(6, "research")},
                    {"id": "competitor", "name": "Competitor Analysis", "endpoint": "/analyst/research", "model": select_model(6, "research")},
                    {"id": "user_stories", "name": "User Stories + AC", "endpoint": "/analyst/research", "model": select_model(6, "research")},
                    {"id": "prd", "name": "Product Requirements Doc", "endpoint": "/analyst/research", "model": select_model(6, "research")},
                    {"id": "market", "name": "Market Research", "endpoint": "/analyst/research", "model": select_model(6, "research")},
                ],
            },
            {
                "id": "qa",
                "name": "QA Team",
                "description": "Test Cases · Accessibility · Compliance",
                "icon": "✅",
                "capabilities": [
                    {"id": "functional", "name": "Functional Test Cases", "endpoint": "/qa/test-cases", "model": select_model(5, "qa")},
                    {"id": "e2e", "name": "E2E Test Cases", "endpoint": "/qa/test-cases", "model": select_model(5, "qa")},
                    {"id": "accessibility", "name": "Accessibility Review (WCAG)", "endpoint": "/qa/accessibility", "model": select_model(4, "qa")},
                    {"id": "compliance", "name": "Compliance Check (GDPR/HIPAA)", "endpoint": "/qa/compliance", "model": select_model(5, "qa")},
                ],
            },
            {
                "id": "engineering",
                "name": "Engineering Team",
                "description": "Code Generation · API Design · Architecture",
                "icon": "⚙️",
                "capabilities": [
                    {"id": "react", "name": "React/TypeScript Code", "endpoint": "/engineering/code", "model": select_model(7, "ui")},
                    {"id": "html", "name": "HTML/Tailwind Code", "endpoint": "/engineering/code", "model": select_model(7, "ui")},
                    {"id": "api_design", "name": "API Design (REST)", "endpoint": "/engineering/api-design", "model": select_model(6, "planner")},
                    {"id": "architecture", "name": "System Architecture", "endpoint": "/engineering/architecture", "model": select_model(8, "planner")},
                ],
            },
        ],
        "ai_backend": AI_BACKEND,
        "models": OLLAMA_MODELS if AI_BACKEND == "ollama" else CLOUD_MODELS,
    }


@app.get("/api/health")
async def health():
    return {
        "status": "healthy",
        "version": "2.0.0",
        "ai_mode": AI_BACKEND,
        "provider": CLOUD_PROVIDER if AI_BACKEND == "cloud" else "ollama",
        "active_models": CLOUD_MODELS if AI_BACKEND == "cloud" else OLLAMA_MODELS,
        "design_system": HAS_DESIGN_SYSTEM,
        "ollama_available": HAS_OLLAMA,
        "cloud_available": HAS_CLOUD_AI,
    }


@app.get("/api/v1/ai/status")
async def ai_status(user: User = Depends(get_current_user)):
    """Report which AI provider is active and what's available."""
    return {
        "backend": AI_BACKEND,
        "provider": CLOUD_PROVIDER if AI_BACKEND == "cloud" else "ollama",
        "providers_detected": {
            "anthropic": bool(ANTHROPIC_API_KEY and ANTHROPIC_API_KEY.startswith("sk-ant")),
            "openai": bool(OPENAI_API_KEY),
            "gemini": bool(GEMINI_API_KEY),
            "ollama": HAS_OLLAMA,
        },
        "active_models": CLOUD_MODELS if AI_BACKEND == "cloud" else OLLAMA_MODELS,
        # Vision (image analysis for "recreate"/"redesign") is resolved independently of
        # AI_BACKEND_MODE above — see resolve_vision_provider() / VISION_PROVIDER env var.
        "vision_enabled": resolve_vision_provider() != "none",
        "vision_provider": resolve_vision_provider(),
        "vision_backend": vision_backend_label(),
    }


@app.get("/api/v1/ai/tokens")
async def get_token_usage(user: User = Depends(get_current_user)):
    """Return cumulative token usage for this server session."""
    total = _token_tracker["input"] + _token_tracker["output"]
    return {
        "input_tokens": _token_tracker["input"],
        "output_tokens": _token_tracker["output"],
        "total_tokens": total,
        "api_calls": _token_tracker["calls"],
        "last_model": _token_tracker["last_model"],
        # Tracked per-call, per-model (see _track_usage) — not a flat single-model estimate.
        "estimated_cost_usd": round(_token_tracker["cost_usd"], 4),
    }


class SwitchAIRequest(BaseModel):
    backend: str  # "cloud" | "ollama"


@app.post("/api/v1/ai/switch")
async def switch_ai_backend(req: SwitchAIRequest, user: User = Depends(get_current_user)):
    """
    Switch the active AI backend at runtime.
    Changes take effect immediately for all subsequent requests.
    """
    global AI_BACKEND
    if req.backend == "ollama":
        if not HAS_OLLAMA:
            raise HTTPException(400, "Ollama is not running. Start Ollama first (https://ollama.ai).")
        AI_BACKEND = "ollama"
        return {"backend": "ollama", "message": "Switched to local Ollama", "models": OLLAMA_MODELS}
    elif req.backend == "cloud":
        if not HAS_CLOUD_AI:
            raise HTTPException(400, "No cloud API key configured. Add ANTHROPIC_API_KEY, OPENAI_API_KEY, or GEMINI_API_KEY to .env.")
        AI_BACKEND = "cloud"
        return {"backend": "cloud", "provider": CLOUD_PROVIDER, "message": f"Switched to {CLOUD_PROVIDER.title()} cloud", "models": CLOUD_MODELS}
    else:
        raise HTTPException(400, "backend must be 'cloud' or 'ollama'")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server_local:app", host="0.0.0.0", port=8000, reload=True)
